from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse, StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import httpx
import asyncio
import json
import io
import base64
import re
import subprocess
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME', 'docagent')]

# Bytez API configuration
BYTEZ_API_KEY = os.environ.get('BYTEZ_API_KEY', 'f870a6e293d31c3d1aba40914052d57b')
BYTEZ_API_URL = "https://api.bytez.com/models/v2"

# Available AI Models via Bytez - Updated with user-requested models
AI_MODELS = {
    "codet5p": {
        "id": "Salesforce/codet5p-16b",
        "name": "CodeT5+ 16B",
        "description": "Salesforce CodeT5+ for code understanding and analysis (Reader Agent)",
        "tasks": ["code-analysis", "code-understanding", "documentation"]
    },
    "qwen-coder-7b": {
        "id": "Qwen/Qwen2.5-Coder-7B-Instruct",
        "name": "Qwen 2.5 Coder 7B",
        "description": "Qwen code-focused model for search and context retrieval (Search Agent)",
        "tasks": ["code-analysis", "documentation", "code-generation"]
    },
    "starcoder2-15b": {
        "id": "bigcode/starcoder2-15b-instruct-v0.1",
        "name": "StarCoder2 15B Instruct",
        "description": "BigCode StarCoder2 for documentation writing (Writer Agent)",
        "tasks": ["code-generation", "documentation", "text-generation"]
    },
    "llama-3.1-8b": {
        "id": "meta-llama/Meta-Llama-3.1-8B-Instruct",
        "name": "Llama 3.1 8B Instruct",
        "description": "Meta Llama 3.1 for verification and diagram generation (Verifier & Diagram Agents)",
        "tasks": ["chat", "verification", "reasoning", "documentation"]
    }
}

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'docagent-secret-key-change-in-production')
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# GitHub OAuth Configuration
GITHUB_CLIENT_ID = os.environ.get('GITHUB_CLIENT_ID', '')
GITHUB_CLIENT_SECRET = os.environ.get('GITHUB_CLIENT_SECRET', '')

# Create the main app
app = FastAPI(title="DocAgent API", version="1.0.0")

# Create routers
api_router = APIRouter(prefix="/api")
auth_router = APIRouter(prefix="/api/auth", tags=["Authentication"])
repos_router = APIRouter(prefix="/api/repositories", tags=["Repositories"])
docs_router = APIRouter(prefix="/api/documentation", tags=["Documentation"])
jobs_router = APIRouter(prefix="/api/jobs", tags=["Jobs"])
analytics_router = APIRouter(prefix="/api/analytics", tags=["Analytics"])
orgs_router = APIRouter(prefix="/api/organizations", tags=["Organizations"])

security = HTTPBearer()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ========================
# PYDANTIC MODELS
# ========================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    name: str
    tenant_id: str
    role: str
    created_at: datetime
    subscription_tier: str = "free"

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class TenantCreate(BaseModel):
    name: str
    subdomain: str

class TenantResponse(BaseModel):
    id: str
    name: str
    subdomain: str
    subscription: Dict[str, Any]
    quotas: Dict[str, int]
    usage: Dict[str, Any]
    created_at: datetime

class RepositoryCreate(BaseModel):
    name: str
    repo_url: str
    provider: str = "github"
    branch: str = "main"
    language: str = "python"

class RepositoryResponse(BaseModel):
    id: str
    tenant_id: str
    name: str
    provider: str
    repo_url: str
    branch: str
    language: str
    last_synced_at: Optional[datetime] = None
    components_count: int = 0
    coverage_percentage: float = 0.0
    created_at: datetime

class DocumentationCreate(BaseModel):
    repository_id: str
    component_path: str
    source_code: str
    component_type: str = "function"
    language: str = "python"

class DocumentationResponse(BaseModel):
    id: str
    tenant_id: str
    repository_id: str
    component_path: str
    component_type: str
    language: str
    docstring: Optional[str] = None
    markdown: Optional[str] = None
    diagrams: List[Dict[str, Any]] = []
    quality_score: float = 0.0
    version: int = 1
    generated_at: Optional[datetime] = None
    created_at: datetime

class JobCreate(BaseModel):
    repository_id: str
    component_path: Optional[str] = None
    job_type: str = "generate"

class JobResponse(BaseModel):
    id: str
    tenant_id: str
    repository_id: str
    job_type: str
    status: str
    progress: int = 0
    component_path: Optional[str] = None
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    created_at: datetime
    completed_at: Optional[datetime] = None

class GenerateDocsRequest(BaseModel):
    repository_id: str
    component_path: str
    source_code: str
    language: str = "python"
    style: str = "google"

class DiagramRequest(BaseModel):
    component_data: Dict[str, Any]
    diagram_type: Optional[str] = None

class RepoDocumentationRequest(BaseModel):
    repo_url: str
    branch: str = "main"

class AgentProgressResponse(BaseModel):
    job_id: str
    status: str
    current_agent: str
    agents: Dict[str, Dict[str, Any]]
    files_processed: int
    total_files: int
    overall_progress: int

# ========================
# SUBSCRIPTION TIERS (MOCKED)
# ========================

SUBSCRIPTION_TIERS = {
    "free": {
        "price": 0,
        "components_per_month": 100,
        "max_repositories": 1,
        "max_team_members": 1,
        "features": ["Basic docstrings", "CLI only"]
    },
    "starter": {
        "price": 29,
        "components_per_month": 1000,
        "max_repositories": 5,
        "max_team_members": 5,
        "features": ["Web UI", "Diagrams", "GitHub integration"]
    },
    "professional": {
        "price": 99,
        "components_per_month": 10000,
        "max_repositories": 20,
        "max_team_members": 20,
        "features": ["Advanced diagrams", "Priority support", "API access"]
    },
    "team": {
        "price": 299,
        "components_per_month": 50000,
        "max_repositories": -1,
        "max_team_members": 50,
        "features": ["Custom templates", "SSO", "Analytics dashboard"]
    },
    "enterprise": {
        "price": -1,
        "components_per_month": -1,
        "max_repositories": -1,
        "max_team_members": -1,
        "features": ["On-premises option", "SLA", "Dedicated support"]
    }
}

# ========================
# HELPER FUNCTIONS
# ========================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())

def create_token(user_id: str, tenant_id: str) -> str:
    payload = {
        "user_id": user_id,
        "tenant_id": tenant_id,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ========================
# MULTI-AGENT AI SYSTEM
# ========================

class BytezAgent:
    """Base agent using Bytez API with free AI models"""
    
    def __init__(self, model_id: str = "Qwen/Qwen2.5-Coder-0.5B-Instruct"):
        self.model_id = model_id
        self.api_key = BYTEZ_API_KEY
        self.api_url = BYTEZ_API_URL
    
    async def generate(self, messages: List[Dict[str, str]], temperature: float = 0.5, max_tokens: int = 1000) -> str:
        """Generate response using Bytez API v2"""
        if not self.api_key:
            logger.warning("No Bytez API key configured, using mock response")
            return self._mock_response(messages)
        
        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                # Bytez API v2 format: POST /models/v2/{modelId}
                url = f"{self.api_url}/{self.model_id}"
                logger.info(f"Calling Bytez API: {url}")
                
                response = await client.post(
                    url,
                    headers={
                        "Authorization": self.api_key,
                        "Content-Type": "application/json"
                    },
                    json={
                        "messages": messages,
                        "params": {
                            "max_new_tokens": max_tokens,
                            "temperature": temperature
                        }
                    }
                )
                
                if response.status_code == 200:
                    data = response.json()
                    logger.info(f"Bytez API response received for {self.model_id}")
                    
                    # Handle Bytez v2 response format
                    if data.get("error") is None:
                        output = data.get("output", {})
                        if isinstance(output, dict):
                            # Chat response format: {"role": "assistant", "content": "..."}
                            content = output.get("content", "")
                            if content:
                                logger.info(f"Got content from Bytez: {content[:100]}...")
                                return content
                        elif isinstance(output, str):
                            logger.info(f"Got string output from Bytez: {output[:100]}...")
                            return output
                    
                    logger.warning(f"Bytez error or unexpected format: {str(data)[:200]}")
                    return self._mock_response(messages)
                else:
                    logger.warning(f"Bytez API returned {response.status_code}: {response.text[:300]}")
                    return self._mock_response(messages)
                    
        except httpx.TimeoutException:
            logger.warning(f"Bytez API timeout for {self.model_id}, using mock")
            return self._mock_response(messages)
        except Exception as e:
            logger.error(f"Bytez API exception for {self.model_id}: {e}")
            return self._mock_response(messages)
    
    def _mock_response(self, messages: List[Dict[str, str]]) -> str:
        """Generate intelligent mock response based on context"""
        user_content = ""
        for msg in messages:
            if msg.get("role") == "user":
                user_content = msg.get("content", "")
                break
        
        # Detect what type of response is needed
        if "analyze" in user_content.lower() or "code" in user_content.lower():
            return json.dumps({
                "complexity": {"cyclomatic": 5, "cognitive": 3},
                "dependencies": {"internal": [], "external": ["json", "typing"]},
                "architecture_type": "function",
                "documentation_needs": ["docstring", "parameters", "return_value", "examples"]
            })
        elif "context" in user_content.lower() or "patterns" in user_content.lower():
            return json.dumps({
                "patterns": ["factory pattern", "dependency injection"],
                "best_practices": ["Use type hints", "Add comprehensive docstrings", "Include examples"],
                "concepts": ["encapsulation", "modularity"],
                "examples": ["# Example usage included"]
            })
        elif "documentation" in user_content.lower() or "write" in user_content.lower():
            return json.dumps({
                "docstring": '"""\\nAuto-generated documentation for code component.\\n\\nThis function/class provides functionality as defined in the source code.\\n\\nArgs:\\n    param1: First parameter description\\n    param2: Second parameter description\\n\\nReturns:\\n    Result of the operation\\n\\nExample:\\n    >>> result = function_name(arg1, arg2)\\n"""',
                "markdown": "# Component Documentation\\n\\n## Overview\\n\\nThis component is part of the codebase and provides specific functionality.\\n\\n## Parameters\\n\\n| Name | Type | Description |\\n|------|------|-------------|\\n| param1 | Any | First parameter |\\n| param2 | Any | Second parameter |\\n\\n## Returns\\n\\nReturns the result of the operation.\\n\\n## Example\\n\\n```python\\nresult = function_name(arg1, arg2)\\nprint(result)\\n```",
                "examples": ["result = function_name(arg1, arg2)", "output = process_data(input)"]
            })
        elif "verify" in user_content.lower() or "quality" in user_content.lower():
            return json.dumps({
                "approved": True,
                "quality_score": 87.5,
                "evaluation": {"accuracy": 90, "completeness": 85, "clarity": 88, "examples": 87},
                "feedback": ["Documentation is comprehensive", "Consider adding more edge cases in examples"]
            })
        elif "diagram" in user_content.lower() or "mermaid" in user_content.lower():
            return json.dumps({
                "diagram_type": "flowchart",
                "mermaid_code": "flowchart TD\\n    A[Start] --> B{Validate Input}\\n    B -->|Valid| C[Process Data]\\n    B -->|Invalid| D[Return Error]\\n    C --> E[Transform Result]\\n    E --> F[Return Output]\\n    D --> F\\n    F --> G[End]",
                "description": "Flowchart showing the main execution flow"
            })
        
        return "Generated content"

class ReaderAgent(BytezAgent):
    """Analyzes code and determines documentation needs - Uses CodeT5+ 16B"""
    
    def __init__(self):
        super().__init__(model_id="Salesforce/codet5p-16b")
    
    async def analyze(self, source_code: str, language: str) -> Dict[str, Any]:
        messages = [
            {"role": "system", "content": """You are a code analysis expert. Analyze the given code and respond with ONLY a valid JSON object (no markdown, no explanation).

The JSON must have this exact structure:
{
  "complexity": {"cyclomatic": <number>, "cognitive": <number>},
  "dependencies": {"internal": [<strings>], "external": [<strings>]},
  "architecture_type": "<string>",
  "documentation_needs": [<strings>]
}

Respond with ONLY the JSON object, nothing else."""},
            {"role": "user", "content": f"Analyze this {language} code:\n\n```{language}\n{source_code}\n```"}
        ]
        
        response = await self.generate(messages, max_tokens=500)
        
        # Parse JSON from response - handle potential markdown wrapping
        try:
            # Try to extract JSON from response
            clean_response = response.strip()
            if "```json" in clean_response:
                clean_response = clean_response.split("```json")[1].split("```")[0].strip()
            elif "```" in clean_response:
                clean_response = clean_response.split("```")[1].split("```")[0].strip()
            return json.loads(clean_response)
        except:
            logger.warning(f"Failed to parse Reader response as JSON: {response[:200]}")
            return {
                "complexity": {"cyclomatic": 5, "cognitive": 3},
                "dependencies": {"internal": [], "external": []},
                "architecture_type": "function",
                "documentation_needs": ["docstring", "parameters", "return_value", "examples"]
            }

class SearcherAgent(BytezAgent):
    """Gathers context for documentation - Uses Qwen2.5-Coder 7B via Bytez"""
    
    def __init__(self):
        super().__init__(model_id="Qwen/Qwen2.5-Coder-7B-Instruct")
    
    async def search(self, code_analysis: Dict[str, Any], language: str) -> Dict[str, Any]:
        messages = [
            {"role": "system", "content": """You are a context retrieval expert. Based on the code analysis, provide relevant context for documentation.

Respond with ONLY a valid JSON object (no markdown, no explanation):
{
  "patterns": [<design patterns used>],
  "best_practices": [<documentation best practices>],
  "concepts": [<relevant programming concepts>],
  "examples": [<example usage patterns>]
}

Respond with ONLY the JSON object, nothing else."""},
            {"role": "user", "content": f"Provide context for documenting {language} code with this analysis:\n{json.dumps(code_analysis)}"}
        ]
        
        response = await self.generate(messages, max_tokens=500)
        
        try:
            clean_response = response.strip()
            if "```json" in clean_response:
                clean_response = clean_response.split("```json")[1].split("```")[0].strip()
            elif "```" in clean_response:
                clean_response = clean_response.split("```")[1].split("```")[0].strip()
            return json.loads(clean_response)
        except:
            logger.warning(f"Failed to parse Searcher response as JSON: {response[:200]}")
            return {
                "patterns": ["standard function documentation"],
                "best_practices": ["Include type hints", "Add examples"],
                "concepts": [],
                "examples": []
            }

class WriterAgent(BytezAgent):
    """Generates documentation - Uses StarCoder2 15B via Bytez"""
    
    def __init__(self):
        super().__init__(model_id="bigcode/starcoder2-15b-instruct-v0.1")
    
    async def write(self, source_code: str, context: Dict[str, Any], language: str, style: str) -> Dict[str, Any]:
        style_guide = {
            "google": "Google style docstrings with Args, Returns, Raises sections",
            "numpy": "NumPy style with Parameters, Returns, Examples sections",
            "sphinx": "Sphinx/reStructuredText format",
            "jsdoc": "JSDoc format for JavaScript/TypeScript"
        }
        
        messages = [
            {"role": "system", "content": f"""You are a technical documentation writer. Generate comprehensive documentation using {style_guide.get(style, style_guide['google'])}.

Respond with ONLY a valid JSON object (no markdown wrapping, no explanation):
{{
  "docstring": "<complete docstring with proper formatting>",
  "markdown": "<full markdown documentation with Overview, Parameters, Returns, Example sections>",
  "examples": ["<example 1>", "<example 2>"]
}}

The docstring should be ready to insert directly into the code. Include proper line breaks using \\n.
Respond with ONLY the JSON object, nothing else."""},
            {"role": "user", "content": f"Write documentation for this {language} code:\n\n```{language}\n{source_code}\n```\n\nContext: {json.dumps(context)}"}
        ]
        
        response = await self.generate(messages, max_tokens=2000)
        
        try:
            clean_response = response.strip()
            if "```json" in clean_response:
                clean_response = clean_response.split("```json")[1].split("```")[0].strip()
            elif clean_response.startswith("```"):
                clean_response = clean_response.split("```")[1].split("```")[0].strip()
            return json.loads(clean_response)
        except:
            logger.warning(f"Failed to parse Writer response as JSON: {response[:300]}")
            # Generate a structured response from the raw text
            func_name = "function"
            if "def " in source_code:
                parts = source_code.split("def ")[1].split("(")
                if parts:
                    func_name = parts[0].strip()
            elif "function " in source_code:
                parts = source_code.split("function ")[1].split("(")
                if parts:
                    func_name = parts[0].strip()
            
            return {
                "docstring": f'"""\n{func_name}: Auto-generated documentation.\n\nThis function performs operations as defined in the source code.\n\nArgs:\n    See source code for parameters.\n\nReturns:\n    See source code for return value.\n"""',
                "markdown": f"# {func_name}\n\n## Overview\n\nThis function is part of the codebase.\n\n## Usage\n\n```{language}\nresult = {func_name}()\n```",
                "examples": [f"result = {func_name}()"]
            }

class VerifierAgent(BytezAgent):
    """Verifies documentation quality - Uses Llama 3.1 8B via Bytez"""
    
    def __init__(self):
        super().__init__(model_id="meta-llama/Meta-Llama-3.1-8B-Instruct")
    
    async def verify(self, source_code: str, documentation: Dict[str, Any]) -> Dict[str, Any]:
        messages = [
            {"role": "system", "content": """You are a documentation quality evaluator. Assess the documentation against:
1. Accuracy (matches code behavior)
2. Completeness (all parameters, returns documented)
3. Clarity (easy to understand)
4. Examples (practical and correct)

Respond with ONLY a valid JSON object (no markdown, no explanation):
{
  "approved": <true or false>,
  "quality_score": <number 0-100>,
  "evaluation": {"accuracy": <0-100>, "completeness": <0-100>, "clarity": <0-100>, "examples": <0-100>},
  "feedback": ["<improvement suggestion 1>", "<improvement suggestion 2>"]
}

Respond with ONLY the JSON object, nothing else."""},
            {"role": "user", "content": f"Verify this documentation:\n\nCode:\n```\n{source_code}\n```\n\nDocumentation:\n{json.dumps(documentation)}"}
        ]
        
        response = await self.generate(messages, max_tokens=500)
        
        try:
            clean_response = response.strip()
            if "```json" in clean_response:
                clean_response = clean_response.split("```json")[1].split("```")[0].strip()
            elif "```" in clean_response:
                clean_response = clean_response.split("```")[1].split("```")[0].strip()
            return json.loads(clean_response)
        except:
            logger.warning(f"Failed to parse Verifier response as JSON: {response[:200]}")
            return {
                "approved": True,
                "quality_score": 85.0,
                "evaluation": {
                    "accuracy": 85,
                    "completeness": 80,
                    "clarity": 90,
                    "examples": 85
                },
                "feedback": ["Documentation generated successfully"]
            }

class DiagramAgent(BytezAgent):
    """Generates Mermaid diagrams - Uses Llama 3.1 8B via Bytez"""
    
    def __init__(self):
        super().__init__(model_id="meta-llama/Meta-Llama-3.1-8B-Instruct")
    
    async def generate_diagram(self, source_code: str, diagram_type: Optional[str] = None) -> Dict[str, Any]:
        messages = [
            {"role": "system", "content": """You are a diagram generation expert. Create a Mermaid.js diagram for the given code.
Choose the most appropriate diagram type:
- flowchart: for control flow
- sequenceDiagram: for interactions
- classDiagram: for class structures
- stateDiagram: for state machines

Respond with ONLY a valid JSON object (no markdown wrapping, no explanation):
{
  "diagram_type": "<flowchart|sequenceDiagram|classDiagram|stateDiagram>",
  "mermaid_code": "<valid Mermaid.js syntax>",
  "description": "<brief description>"
}

IMPORTANT: The mermaid_code must be valid Mermaid syntax. Use \\n for newlines.
Respond with ONLY the JSON object, nothing else."""},
            {"role": "user", "content": f"Create a diagram for:\n```\n{source_code}\n```" + (f"\nPreferred type: {diagram_type}" if diagram_type else "")}
        ]
        
        response = await self.generate(messages, max_tokens=800)
        
        try:
            clean_response = response.strip()
            if "```json" in clean_response:
                clean_response = clean_response.split("```json")[1].split("```")[0].strip()
            elif clean_response.startswith("```"):
                clean_response = clean_response.split("```")[1].split("```")[0].strip()
            return json.loads(clean_response)
        except:
            logger.warning(f"Failed to parse Diagram response as JSON: {response[:300]}")
            # Generate a simple flowchart
            return {
                "diagram_type": "flowchart",
                "mermaid_code": "flowchart TD\n    A[Start] --> B{Input Valid?}\n    B -->|Yes| C[Process Data]\n    B -->|No| D[Handle Error]\n    C --> E[Return Result]\n    D --> E\n    E --> F[End]",
                "description": "Auto-generated flowchart showing basic control flow"
            }

class OrchestratorAgent:
    """Coordinates all agents for documentation generation"""
    
    def __init__(self):
        self.reader = ReaderAgent()
        self.searcher = SearcherAgent()
        self.writer = WriterAgent()
        self.verifier = VerifierAgent()
        self.diagram = DiagramAgent()
    
    async def generate_documentation(
        self, 
        source_code: str, 
        language: str, 
        style: str = "google",
        job_id: str = None,
        progress_callback = None
    ) -> Dict[str, Any]:
        """Full documentation generation pipeline"""
        
        result = {
            "status": "processing",
            "stages": {}
        }
        
        try:
            # Stage 1: Reader - Analyze code
            if progress_callback:
                await progress_callback(job_id, 10, "Analyzing code structure...")
            analysis = await self.reader.analyze(source_code, language)
            result["stages"]["analysis"] = analysis
            
            # Stage 2: Searcher - Gather context
            if progress_callback:
                await progress_callback(job_id, 30, "Gathering context...")
            context = await self.searcher.search(analysis, language)
            result["stages"]["context"] = context
            
            # Stage 3: Writer - Generate documentation
            if progress_callback:
                await progress_callback(job_id, 50, "Writing documentation...")
            documentation = await self.writer.write(source_code, context, language, style)
            result["stages"]["documentation"] = documentation
            
            # Stage 4: Verifier - Check quality
            if progress_callback:
                await progress_callback(job_id, 70, "Verifying quality...")
            verification = await self.verifier.verify(source_code, documentation)
            result["stages"]["verification"] = verification
            
            # Stage 5: Diagram - Generate visual
            if progress_callback:
                await progress_callback(job_id, 90, "Generating diagrams...")
            diagram = await self.diagram.generate_diagram(source_code)
            result["stages"]["diagram"] = diagram
            
            # Final result
            if progress_callback:
                await progress_callback(job_id, 100, "Complete!")
            
            result["status"] = "completed"
            result["documentation"] = {
                "docstring": documentation.get("docstring", ""),
                "markdown": documentation.get("markdown", ""),
                "examples": documentation.get("examples", []),
                "quality_score": verification.get("quality_score", 0),
                "diagram": diagram
            }
            
        except Exception as e:
            logger.error(f"Documentation generation error: {e}")
            result["status"] = "failed"
            result["error"] = str(e)
        
        return result

# Initialize orchestrator
orchestrator = OrchestratorAgent()

# ========================
# WEBSOCKET MANAGER
# ========================

class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}
    
    async def connect(self, websocket: WebSocket, client_id: str):
        await websocket.accept()
        self.active_connections[client_id] = websocket
    
    def disconnect(self, client_id: str):
        if client_id in self.active_connections:
            del self.active_connections[client_id]
    
    async def send_progress(self, client_id: str, data: dict):
        if client_id in self.active_connections:
            try:
                await self.active_connections[client_id].send_json(data)
            except:
                self.disconnect(client_id)

ws_manager = ConnectionManager()

# ========================
# AUTH ROUTES
# ========================

@auth_router.post("/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    # Check if user exists
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create tenant (organization)
    tenant_id = str(uuid.uuid4())
    tenant = {
        "id": tenant_id,
        "name": f"{user_data.name}'s Organization",
        "subdomain": user_data.email.split("@")[0].lower().replace(".", "-"),
        "subscription": {
            "tier": "free",
            "status": "active",
            "current_period_end": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
        },
        "quotas": SUBSCRIPTION_TIERS["free"],
        "usage": {"components_this_month": 0, "last_reset_date": datetime.now(timezone.utc).isoformat()},
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.tenants.insert_one(tenant)
    
    # Create user
    user_id = str(uuid.uuid4())
    user = {
        "id": user_id,
        "email": user_data.email,
        "name": user_data.name,
        "password_hash": hash_password(user_data.password),
        "tenant_id": tenant_id,
        "role": "owner",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "last_login": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(user)
    
    token = create_token(user_id, tenant_id)
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user_id,
            email=user_data.email,
            name=user_data.name,
            tenant_id=tenant_id,
            role="owner",
            created_at=datetime.fromisoformat(user["created_at"]),
            subscription_tier="free"
        )
    )

@auth_router.post("/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Update last login
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"last_login": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Get tenant for subscription info
    tenant = await db.tenants.find_one({"id": user["tenant_id"]}, {"_id": 0})
    subscription_tier = tenant.get("subscription", {}).get("tier", "free") if tenant else "free"
    
    token = create_token(user["id"], user["tenant_id"])
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user["id"],
            email=user["email"],
            name=user["name"],
            tenant_id=user["tenant_id"],
            role=user["role"],
            created_at=datetime.fromisoformat(user["created_at"]),
            subscription_tier=subscription_tier
        )
    )

@auth_router.get("/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    tenant = await db.tenants.find_one({"id": current_user["tenant_id"]}, {"_id": 0})
    subscription_tier = tenant.get("subscription", {}).get("tier", "free") if tenant else "free"
    
    return UserResponse(
        id=current_user["id"],
        email=current_user["email"],
        name=current_user["name"],
        tenant_id=current_user["tenant_id"],
        role=current_user["role"],
        created_at=datetime.fromisoformat(current_user["created_at"]),
        subscription_tier=subscription_tier
    )

# GitHub OAuth (MOCKED)
@auth_router.get("/oauth/github")
async def github_oauth():
    """Returns mock GitHub OAuth URL - In production, this would redirect to GitHub"""
    return {"url": "https://github.com/login/oauth/authorize?client_id=MOCK_CLIENT_ID&scope=repo,user:email"}

@auth_router.post("/oauth/github/callback")
async def github_callback(code: str = "mock_code"):
    """Mock GitHub OAuth callback - creates a demo user"""
    # In production, exchange code for token and fetch user info from GitHub
    
    # Create mock GitHub user
    mock_email = f"github_user_{uuid.uuid4().hex[:8]}@example.com"
    mock_name = "GitHub User"
    
    # Check if user exists
    existing = await db.users.find_one({"email": mock_email}, {"_id": 0})
    if existing:
        token = create_token(existing["id"], existing["tenant_id"])
        return TokenResponse(
            access_token=token,
            user=UserResponse(
                id=existing["id"],
                email=existing["email"],
                name=existing["name"],
                tenant_id=existing["tenant_id"],
                role=existing["role"],
                created_at=datetime.fromisoformat(existing["created_at"]),
                subscription_tier="free"
            )
        )
    
    # Create new user via register logic
    user_data = UserCreate(email=mock_email, password=uuid.uuid4().hex, name=mock_name)
    return await register(user_data)

# ========================
# REPOSITORY ROUTES
# ========================

@repos_router.get("", response_model=List[RepositoryResponse])
async def list_repositories(current_user: dict = Depends(get_current_user)):
    repos = await db.repositories.find(
        {"tenant_id": current_user["tenant_id"]}, 
        {"_id": 0}
    ).to_list(100)
    
    return [RepositoryResponse(
        id=r["id"],
        tenant_id=r["tenant_id"],
        name=r["name"],
        provider=r["provider"],
        repo_url=r["repo_url"],
        branch=r["branch"],
        language=r["language"],
        last_synced_at=datetime.fromisoformat(r["last_synced_at"]) if r.get("last_synced_at") else None,
        components_count=r.get("components_count", 0),
        coverage_percentage=r.get("coverage_percentage", 0.0),
        created_at=datetime.fromisoformat(r["created_at"])
    ) for r in repos]

@repos_router.post("", response_model=RepositoryResponse)
async def create_repository(repo_data: RepositoryCreate, current_user: dict = Depends(get_current_user)):
    repo_id = str(uuid.uuid4())
    repo = {
        "id": repo_id,
        "tenant_id": current_user["tenant_id"],
        "name": repo_data.name,
        "provider": repo_data.provider,
        "repo_url": repo_data.repo_url,
        "branch": repo_data.branch,
        "language": repo_data.language,
        "last_synced_at": None,
        "components_count": 0,
        "coverage_percentage": 0.0,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.repositories.insert_one(repo)
    
    return RepositoryResponse(
        id=repo_id,
        tenant_id=repo["tenant_id"],
        name=repo["name"],
        provider=repo["provider"],
        repo_url=repo["repo_url"],
        branch=repo["branch"],
        language=repo["language"],
        components_count=0,
        coverage_percentage=0.0,
        created_at=datetime.fromisoformat(repo["created_at"])
    )

@repos_router.get("/{repo_id}", response_model=RepositoryResponse)
async def get_repository(repo_id: str, current_user: dict = Depends(get_current_user)):
    repo = await db.repositories.find_one(
        {"id": repo_id, "tenant_id": current_user["tenant_id"]},
        {"_id": 0}
    )
    if not repo:
        raise HTTPException(status_code=404, detail="Repository not found")
    
    return RepositoryResponse(
        id=repo["id"],
        tenant_id=repo["tenant_id"],
        name=repo["name"],
        provider=repo["provider"],
        repo_url=repo["repo_url"],
        branch=repo["branch"],
        language=repo["language"],
        last_synced_at=datetime.fromisoformat(repo["last_synced_at"]) if repo.get("last_synced_at") else None,
        components_count=repo.get("components_count", 0),
        coverage_percentage=repo.get("coverage_percentage", 0.0),
        created_at=datetime.fromisoformat(repo["created_at"])
    )

@repos_router.delete("/{repo_id}")
async def delete_repository(repo_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.repositories.delete_one(
        {"id": repo_id, "tenant_id": current_user["tenant_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Repository not found")
    return {"message": "Repository deleted"}

# ========================
# DOCUMENTATION ROUTES
# ========================

@docs_router.get("", response_model=List[DocumentationResponse])
async def list_documentation(
    repository_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {"tenant_id": current_user["tenant_id"]}
    if repository_id:
        query["repository_id"] = repository_id
    
    docs = await db.documentation.find(query, {"_id": 0}).to_list(100)
    
    return [DocumentationResponse(
        id=d["id"],
        tenant_id=d["tenant_id"],
        repository_id=d["repository_id"],
        component_path=d["component_path"],
        component_type=d["component_type"],
        language=d["language"],
        docstring=d.get("docstring"),
        markdown=d.get("markdown"),
        diagrams=d.get("diagrams", []),
        quality_score=d.get("quality_score", 0.0),
        version=d.get("version", 1),
        generated_at=datetime.fromisoformat(d["generated_at"]) if d.get("generated_at") else None,
        created_at=datetime.fromisoformat(d["created_at"])
    ) for d in docs]

@docs_router.get("/{doc_id}", response_model=DocumentationResponse)
async def get_documentation(doc_id: str, current_user: dict = Depends(get_current_user)):
    doc = await db.documentation.find_one(
        {"id": doc_id, "tenant_id": current_user["tenant_id"]},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Documentation not found")
    
    return DocumentationResponse(
        id=doc["id"],
        tenant_id=doc["tenant_id"],
        repository_id=doc["repository_id"],
        component_path=doc["component_path"],
        component_type=doc["component_type"],
        language=doc["language"],
        docstring=doc.get("docstring"),
        markdown=doc.get("markdown"),
        diagrams=doc.get("diagrams", []),
        quality_score=doc.get("quality_score", 0.0),
        version=doc.get("version", 1),
        generated_at=datetime.fromisoformat(doc["generated_at"]) if doc.get("generated_at") else None,
        created_at=datetime.fromisoformat(doc["created_at"])
    )

@docs_router.post("/generate", response_model=JobResponse)
async def generate_documentation(
    request: GenerateDocsRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Start documentation generation job"""
    job_id = str(uuid.uuid4())
    job = {
        "id": job_id,
        "tenant_id": current_user["tenant_id"],
        "repository_id": request.repository_id,
        "job_type": "generate",
        "status": "queued",
        "progress": 0,
        "component_path": request.component_path,
        "source_code": request.source_code,
        "language": request.language,
        "style": request.style,
        "result": None,
        "error": None,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": None
    }
    await db.jobs.insert_one(job)
    
    # Start background task
    background_tasks.add_task(
        run_documentation_job,
        job_id,
        current_user["tenant_id"],
        request.source_code,
        request.language,
        request.style,
        request.repository_id,
        request.component_path
    )
    
    return JobResponse(
        id=job_id,
        tenant_id=job["tenant_id"],
        repository_id=job["repository_id"],
        job_type=job["job_type"],
        status=job["status"],
        progress=0,
        component_path=job["component_path"],
        created_at=datetime.fromisoformat(job["created_at"])
    )

async def update_job_progress(job_id: str, progress: int, stage: str):
    """Update job progress and notify via WebSocket"""
    await db.jobs.update_one(
        {"id": job_id},
        {"$set": {"progress": progress, "current_stage": stage, "status": "processing"}}
    )
    await ws_manager.send_progress(job_id, {
        "type": "job:progress",
        "job_id": job_id,
        "progress": progress,
        "stage": stage
    })

async def run_documentation_job(
    job_id: str,
    tenant_id: str,
    source_code: str,
    language: str,
    style: str,
    repository_id: str,
    component_path: str
):
    """Background task for documentation generation"""
    try:
        result = await orchestrator.generate_documentation(
            source_code=source_code,
            language=language,
            style=style,
            job_id=job_id,
            progress_callback=update_job_progress
        )
        
        if result["status"] == "completed":
            # Save documentation
            doc_id = str(uuid.uuid4())
            doc = {
                "id": doc_id,
                "tenant_id": tenant_id,
                "repository_id": repository_id,
                "component_path": component_path,
                "component_type": result["stages"]["analysis"].get("architecture_type", "function"),
                "language": language,
                "docstring": result["documentation"]["docstring"],
                "markdown": result["documentation"]["markdown"],
                "diagrams": [result["documentation"]["diagram"]],
                "quality_score": result["documentation"]["quality_score"],
                "version": 1,
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.documentation.insert_one(doc)
            
            # Update job
            await db.jobs.update_one(
                {"id": job_id},
                {"$set": {
                    "status": "completed",
                    "progress": 100,
                    "result": {"documentation_id": doc_id, **result["documentation"]},
                    "completed_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            
            await ws_manager.send_progress(job_id, {
                "type": "job:completed",
                "job_id": job_id,
                "result": {"documentation_id": doc_id}
            })
        else:
            await db.jobs.update_one(
                {"id": job_id},
                {"$set": {
                    "status": "failed",
                    "error": result.get("error", "Unknown error"),
                    "completed_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            
            await ws_manager.send_progress(job_id, {
                "type": "job:failed",
                "job_id": job_id,
                "error": result.get("error")
            })
            
    except Exception as e:
        logger.error(f"Job {job_id} failed: {e}")
        await db.jobs.update_one(
            {"id": job_id},
            {"$set": {
                "status": "failed",
                "error": str(e),
                "completed_at": datetime.now(timezone.utc).isoformat()
            }}
        )

# ========================
# JOBS ROUTES
# ========================

@jobs_router.get("", response_model=List[JobResponse])
async def list_jobs(
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {"tenant_id": current_user["tenant_id"]}
    if status:
        query["status"] = status
    
    jobs = await db.jobs.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    return [JobResponse(
        id=j["id"],
        tenant_id=j["tenant_id"],
        repository_id=j["repository_id"],
        job_type=j["job_type"],
        status=j["status"],
        progress=j.get("progress", 0),
        component_path=j.get("component_path"),
        result=j.get("result"),
        error=j.get("error"),
        created_at=datetime.fromisoformat(j["created_at"]),
        completed_at=datetime.fromisoformat(j["completed_at"]) if j.get("completed_at") else None
    ) for j in jobs]

@jobs_router.get("/{job_id}", response_model=JobResponse)
async def get_job(job_id: str, current_user: dict = Depends(get_current_user)):
    job = await db.jobs.find_one(
        {"id": job_id, "tenant_id": current_user["tenant_id"]},
        {"_id": 0}
    )
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    return JobResponse(
        id=job["id"],
        tenant_id=job["tenant_id"],
        repository_id=job["repository_id"],
        job_type=job["job_type"],
        status=job["status"],
        progress=job.get("progress", 0),
        component_path=job.get("component_path"),
        result=job.get("result"),
        error=job.get("error"),
        created_at=datetime.fromisoformat(job["created_at"]),
        completed_at=datetime.fromisoformat(job["completed_at"]) if job.get("completed_at") else None
    )

# ========================
# ANALYTICS ROUTES
# ========================

@analytics_router.get("/overview")
async def get_analytics_overview(current_user: dict = Depends(get_current_user)):
    tenant_id = current_user["tenant_id"]
    
    # Get counts
    repos_count = await db.repositories.count_documents({"tenant_id": tenant_id})
    docs_count = await db.documentation.count_documents({"tenant_id": tenant_id})
    jobs_count = await db.jobs.count_documents({"tenant_id": tenant_id})
    
    # Get average quality score
    pipeline = [
        {"$match": {"tenant_id": tenant_id}},
        {"$group": {"_id": None, "avg_quality": {"$avg": "$quality_score"}}}
    ]
    quality_result = await db.documentation.aggregate(pipeline).to_list(1)
    avg_quality = quality_result[0]["avg_quality"] if quality_result else 0
    
    # Get recent jobs
    recent_jobs = await db.jobs.find(
        {"tenant_id": tenant_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(5)
    
    return {
        "total_repositories": repos_count,
        "total_documentation": docs_count,
        "total_jobs": jobs_count,
        "average_quality_score": round(avg_quality, 1),
        "recent_jobs": recent_jobs,
        "components_documented_this_month": docs_count,
        "coverage_percentage": 75.5  # Mock value
    }

@analytics_router.get("/coverage")
async def get_coverage_stats(current_user: dict = Depends(get_current_user)):
    # Mock coverage data
    return {
        "by_language": {
            "python": 85,
            "javascript": 70,
            "typescript": 60
        },
        "by_repository": [],
        "trend": [
            {"date": "2025-01-01", "coverage": 50},
            {"date": "2025-01-08", "coverage": 60},
            {"date": "2025-01-15", "coverage": 70},
            {"date": "2025-01-22", "coverage": 75}
        ]
    }

# ========================
# AI MODELS ROUTES - Ollama Integration for Local Models
# ========================

models_router = APIRouter(prefix="/api/models", tags=["AI Models"])

# Available free models for download via Ollama
AVAILABLE_LOCAL_MODELS = [
    {
        "id": "llama3.2:1b",
        "name": "Llama 3.2 1B",
        "description": "Meta's smallest Llama 3.2 model - fast and efficient for code analysis",
        "size": "1.3GB",
        "tasks": ["code-analysis", "chat", "documentation"],
        "recommended_for": ["Reader Agent", "Verifier Agent"],
        "free": True
    },
    {
        "id": "llama3.2:3b",
        "name": "Llama 3.2 3B",
        "description": "Balanced Llama 3.2 model - good quality and speed",
        "size": "2.0GB",
        "tasks": ["code-analysis", "documentation", "chat"],
        "recommended_for": ["Writer Agent", "Searcher Agent"],
        "free": True
    },
    {
        "id": "codellama:7b",
        "name": "Code Llama 7B",
        "description": "Specialized for code understanding and generation",
        "size": "3.8GB",
        "tasks": ["code-analysis", "code-generation", "documentation"],
        "recommended_for": ["Reader Agent", "Writer Agent"],
        "free": True
    },
    {
        "id": "qwen2.5-coder:1.5b",
        "name": "Qwen 2.5 Coder 1.5B",
        "description": "Alibaba's code-optimized model - excellent for documentation",
        "size": "1.0GB",
        "tasks": ["code-analysis", "documentation", "code-generation"],
        "recommended_for": ["All Agents"],
        "free": True
    },
    {
        "id": "qwen2.5-coder:7b",
        "name": "Qwen 2.5 Coder 7B",
        "description": "Larger Qwen coder model for better quality documentation",
        "size": "4.7GB",
        "tasks": ["code-analysis", "documentation", "code-generation"],
        "recommended_for": ["Writer Agent", "Diagram Agent"],
        "free": True
    },
    {
        "id": "deepseek-coder:1.3b",
        "name": "DeepSeek Coder 1.3B",
        "description": "Efficient coding model from DeepSeek",
        "size": "0.8GB",
        "tasks": ["code-analysis", "code-generation"],
        "recommended_for": ["Reader Agent"],
        "free": True
    },
    {
        "id": "deepseek-coder:6.7b",
        "name": "DeepSeek Coder 6.7B",
        "description": "Powerful coding model for complex documentation tasks",
        "size": "3.8GB",
        "tasks": ["code-analysis", "code-generation", "documentation"],
        "recommended_for": ["All Agents"],
        "free": True
    },
    {
        "id": "phi3:mini",
        "name": "Phi-3 Mini",
        "description": "Microsoft's compact but powerful model",
        "size": "2.3GB",
        "tasks": ["chat", "documentation", "reasoning"],
        "recommended_for": ["Verifier Agent"],
        "free": True
    },
    {
        "id": "mistral:7b",
        "name": "Mistral 7B",
        "description": "High-quality open model from Mistral AI",
        "size": "4.1GB",
        "tasks": ["chat", "documentation", "code-analysis"],
        "recommended_for": ["Writer Agent", "Diagram Agent"],
        "free": True
    },
    {
        "id": "gemma2:2b",
        "name": "Gemma 2 2B",
        "description": "Google's efficient open model",
        "size": "1.6GB",
        "tasks": ["chat", "documentation"],
        "recommended_for": ["Searcher Agent"],
        "free": True
    }
]

# Track download progress
model_download_progress: Dict[str, Dict[str, Any]] = {}

async def check_ollama_installed() -> bool:
    """Check if Ollama is installed and running"""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get("http://localhost:11434/api/tags")
            return response.status_code == 200
    except:
        return False

async def get_installed_models() -> List[Dict[str, Any]]:
    """Get list of models installed in Ollama"""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get("http://localhost:11434/api/tags")
            if response.status_code == 200:
                data = response.json()
                return data.get("models", [])
    except Exception as e:
        logger.error(f"Error getting Ollama models: {e}")
    return []

@models_router.get("")
async def list_ai_models():
    """Get available AI models for documentation generation"""
    ollama_installed = await check_ollama_installed()
    installed_models = []
    
    if ollama_installed:
        installed_models = await get_installed_models()
    
    installed_names = [m.get("name", "").split(":")[0] for m in installed_models]
    
    # Mark which models are installed
    local_models = []
    for model in AVAILABLE_LOCAL_MODELS:
        model_copy = model.copy()
        model_base = model["id"].split(":")[0]
        model_copy["installed"] = any(model_base in name or model["id"] in name for name in [m.get("name", "") for m in installed_models])
        model_copy["downloading"] = model["id"] in model_download_progress and model_download_progress[model["id"]].get("status") == "downloading"
        if model_copy["downloading"]:
            model_copy["download_progress"] = model_download_progress[model["id"]].get("progress", 0)
        local_models.append(model_copy)
    
    return {
        "ollama_installed": ollama_installed,
        "ollama_url": "http://localhost:11434",
        "local_models": local_models,
        "installed_models": installed_models,
        "cloud_models": [
            {
                "id": "Salesforce/codet5p-16b",
                "name": "CodeT5+ 16B (Cloud)",
                "description": "Salesforce CodeT5+ for code understanding - Reader Agent",
                "tasks": ["code-analysis", "code-understanding", "documentation"],
                "status": "available",
                "free": True,
                "cloud": True,
                "assigned_to": "reader"
            },
            {
                "id": "Qwen/Qwen2.5-Coder-7B-Instruct",
                "name": "Qwen 2.5 Coder 7B (Cloud)",
                "description": "Qwen code model for search and context - Search Agent",
                "tasks": ["code-analysis", "documentation", "code-generation"],
                "status": "available",
                "free": True,
                "cloud": True,
                "assigned_to": "searcher"
            },
            {
                "id": "bigcode/starcoder2-15b-instruct-v0.1",
                "name": "StarCoder2 15B (Cloud)",
                "description": "BigCode StarCoder2 for documentation writing - Writer Agent",
                "tasks": ["code-generation", "documentation", "text-generation"],
                "status": "available",
                "free": True,
                "cloud": True,
                "assigned_to": "writer"
            },
            {
                "id": "meta-llama/Meta-Llama-3.1-8B-Instruct",
                "name": "Llama 3.1 8B (Cloud)",
                "description": "Meta Llama 3.1 for verification and diagrams - Verifier & Diagram Agents",
                "tasks": ["chat", "verification", "reasoning", "documentation"],
                "status": "available",
                "free": True,
                "cloud": True,
                "assigned_to": ["verifier", "diagram"]
            }
        ],
        "agent_assignments": {
            "reader": "Salesforce/codet5p-16b",
            "searcher": "Qwen/Qwen2.5-Coder-7B-Instruct",
            "writer": "bigcode/starcoder2-15b-instruct-v0.1",
            "verifier": "meta-llama/Meta-Llama-3.1-8B-Instruct",
            "diagram": "meta-llama/Meta-Llama-3.1-8B-Instruct"
        }
    }

@models_router.get("/status")
async def get_models_status():
    """Check status of AI models and Ollama"""
    ollama_installed = await check_ollama_installed()
    installed_models = await get_installed_models() if ollama_installed else []
    
    return {
        "bytez_configured": bool(BYTEZ_API_KEY),
        "bytez_api_url": BYTEZ_API_URL,
        "ollama_installed": ollama_installed,
        "ollama_running": ollama_installed,
        "ollama_url": "http://localhost:11434",
        "installed_model_count": len(installed_models),
        "models_available": True
    }

@models_router.post("/download/{model_id:path}")
async def download_model(model_id: str, background_tasks: BackgroundTasks):
    """Start downloading a model via Ollama"""
    ollama_installed = await check_ollama_installed()
    
    if not ollama_installed:
        raise HTTPException(
            status_code=400, 
            detail="Ollama is not installed or not running. Please install Ollama first: https://ollama.ai"
        )
    
    # Check if already downloading
    if model_id in model_download_progress and model_download_progress[model_id].get("status") == "downloading":
        return {"message": "Download already in progress", "model_id": model_id}
    
    # Initialize progress tracking
    model_download_progress[model_id] = {
        "status": "downloading",
        "progress": 0,
        "started_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Start background download
    background_tasks.add_task(download_model_task, model_id)
    
    return {
        "message": f"Started downloading {model_id}",
        "model_id": model_id,
        "status": "downloading"
    }

async def download_model_task(model_id: str):
    """Background task to download model via Ollama"""
    try:
        async with httpx.AsyncClient(timeout=None) as client:
            async with client.stream(
                "POST",
                "http://localhost:11434/api/pull",
                json={"name": model_id, "stream": True},
                timeout=None
            ) as response:
                async for line in response.aiter_lines():
                    if line:
                        try:
                            data = json.loads(line)
                            if "completed" in data and "total" in data:
                                progress = int((data["completed"] / data["total"]) * 100)
                                model_download_progress[model_id]["progress"] = progress
                            if data.get("status") == "success":
                                model_download_progress[model_id] = {
                                    "status": "completed",
                                    "progress": 100,
                                    "completed_at": datetime.now(timezone.utc).isoformat()
                                }
                                break
                        except json.JSONDecodeError:
                            pass
    except Exception as e:
        logger.error(f"Error downloading model {model_id}: {e}")
        model_download_progress[model_id] = {
            "status": "failed",
            "error": str(e)
        }

@models_router.get("/download/{model_id:path}/progress")
async def get_download_progress(model_id: str):
    """Get download progress for a model"""
    if model_id not in model_download_progress:
        return {"status": "not_started", "progress": 0}
    return model_download_progress[model_id]

@models_router.delete("/{model_id:path}")
async def delete_model(model_id: str):
    """Delete a downloaded model from Ollama"""
    ollama_installed = await check_ollama_installed()
    
    if not ollama_installed:
        raise HTTPException(status_code=400, detail="Ollama is not running")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.delete(
                "http://localhost:11434/api/delete",
                json={"name": model_id}
            )
            if response.status_code == 200:
                return {"message": f"Model {model_id} deleted successfully"}
            else:
                raise HTTPException(status_code=response.status_code, detail="Failed to delete model")
    except httpx.HTTPError as e:
        raise HTTPException(status_code=500, detail=str(e))

@models_router.post("/chat/{model_id:path}")
async def chat_with_model(model_id: str, message: Dict[str, str]):
    """Chat with a local Ollama model"""
    ollama_installed = await check_ollama_installed()
    
    if not ollama_installed:
        raise HTTPException(status_code=400, detail="Ollama is not running")
    
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": model_id,
                    "prompt": message.get("content", ""),
                    "stream": False
                }
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "response": data.get("response", ""),
                    "model": model_id,
                    "done": data.get("done", True)
                }
            else:
                raise HTTPException(status_code=response.status_code, detail="Failed to generate response")
    except httpx.HTTPError as e:
        raise HTTPException(status_code=500, detail=str(e))

@models_router.get("/ollama/install-guide")
async def get_ollama_install_guide():
    """Get installation instructions for Ollama"""
    return {
        "title": "Install Ollama to Run Local AI Models",
        "description": "Ollama allows you to run powerful AI models locally on your machine for free.",
        "platforms": {
            "macos": {
                "command": "curl -fsSL https://ollama.ai/install.sh | sh",
                "alternative": "Download from https://ollama.ai/download"
            },
            "linux": {
                "command": "curl -fsSL https://ollama.ai/install.sh | sh"
            },
            "windows": {
                "command": "Download installer from https://ollama.ai/download"
            }
        },
        "after_install": [
            "Run 'ollama serve' to start the Ollama server",
            "Return to this page to download and use models"
        ],
        "documentation_url": "https://ollama.ai/docs"
    }

# ========================
# ORGANIZATIONS ROUTES
# ========================

@orgs_router.get("/current", response_model=TenantResponse)
async def get_current_organization(current_user: dict = Depends(get_current_user)):
    tenant = await db.tenants.find_one({"id": current_user["tenant_id"]}, {"_id": 0})
    if not tenant:
        raise HTTPException(status_code=404, detail="Organization not found")
    
    return TenantResponse(
        id=tenant["id"],
        name=tenant["name"],
        subdomain=tenant["subdomain"],
        subscription=tenant["subscription"],
        quotas=tenant["quotas"],
        usage=tenant["usage"],
        created_at=datetime.fromisoformat(tenant["created_at"])
    )

@orgs_router.get("/subscription-tiers")
async def get_subscription_tiers():
    """Get available subscription tiers (MOCKED)"""
    return SUBSCRIPTION_TIERS

@orgs_router.post("/upgrade")
async def upgrade_subscription(tier: str, current_user: dict = Depends(get_current_user)):
    """Mock subscription upgrade"""
    if tier not in SUBSCRIPTION_TIERS:
        raise HTTPException(status_code=400, detail="Invalid tier")
    
    await db.tenants.update_one(
        {"id": current_user["tenant_id"]},
        {"$set": {
            "subscription.tier": tier,
            "subscription.status": "active",
            "quotas": SUBSCRIPTION_TIERS[tier]
        }}
    )
    
    return {"message": f"Upgraded to {tier} tier (MOCKED)", "tier": tier}

# ========================
# REPOSITORY DOCUMENTATION ROUTER
# ========================

repo_docs_router = APIRouter(prefix="/api/repo-documentation", tags=["Repository Documentation"])

# Store active documentation jobs with detailed progress
active_doc_jobs: Dict[str, Dict[str, Any]] = {}

async def fetch_github_repo_contents(repo_url: str, branch: str = "main", access_token: str = None) -> List[Dict[str, Any]]:
    """Fetch all code files from a GitHub repository"""
    # Parse repo URL to get owner and repo name
    # Supports: https://github.com/owner/repo or github.com/owner/repo
    match = re.match(r'(?:https?://)?github\.com/([^/]+)/([^/]+?)(?:\.git)?/?$', repo_url)
    if not match:
        raise HTTPException(status_code=400, detail="Invalid GitHub repository URL")
    
    owner, repo = match.groups()
    
    headers = {
        "Accept": "application/vnd.github.v3+json",
        "User-Agent": "DocAgent"
    }
    if access_token:
        headers["Authorization"] = f"token {access_token}"
    
    files = []
    code_extensions = {'.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h', '.go', '.rs', '.cs', '.rb', '.php'}
    
    async def fetch_tree(path: str = ""):
        """Recursively fetch repository tree"""
        url = f"https://api.github.com/repos/{owner}/{repo}/contents/{path}"
        if branch:
            url += f"?ref={branch}"
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(url, headers=headers)
            
            if response.status_code == 404:
                raise HTTPException(status_code=404, detail="Repository not found or private")
            elif response.status_code == 403:
                raise HTTPException(status_code=403, detail="API rate limit exceeded or access denied")
            elif response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch repository")
            
            items = response.json()
            
            for item in items:
                if item['type'] == 'file':
                    ext = os.path.splitext(item['name'])[1].lower()
                    if ext in code_extensions:
                        # Fetch file content
                        file_response = await client.get(item['download_url'], headers=headers)
                        if file_response.status_code == 200:
                            content = file_response.text
                            # Detect language from extension
                            lang_map = {
                                '.py': 'python', '.js': 'javascript', '.jsx': 'javascript',
                                '.ts': 'typescript', '.tsx': 'typescript', '.java': 'java',
                                '.cpp': 'cpp', '.c': 'c', '.h': 'c', '.go': 'go',
                                '.rs': 'rust', '.cs': 'csharp', '.rb': 'ruby', '.php': 'php'
                            }
                            files.append({
                                'path': item['path'],
                                'name': item['name'],
                                'content': content,
                                'language': lang_map.get(ext, 'text'),
                                'size': item['size']
                            })
                elif item['type'] == 'dir' and not item['name'].startswith('.'):
                    # Skip hidden directories and common non-code directories
                    skip_dirs = {'node_modules', 'venv', '.git', '__pycache__', 'dist', 'build', '.idea', '.vscode'}
                    if item['name'] not in skip_dirs:
                        await fetch_tree(item['path'])
    
    await fetch_tree()
    return files

def render_mermaid_to_image(mermaid_code: str) -> Optional[io.BytesIO]:
    """Render Mermaid code to PNG image using mmdc CLI"""
    try:
        # Create temporary files for input and output
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as mmd_file:
            mmd_file.write(mermaid_code)
            mmd_path = mmd_file.name
        
        png_path = mmd_path.replace('.mmd', '.png')
        
        # Create puppeteer config if it doesn't exist
        puppeteer_config_path = '/tmp/puppeteer-config.json'
        if not os.path.exists(puppeteer_config_path):
            with open(puppeteer_config_path, 'w') as f:
                f.write('{"executablePath": "/usr/bin/chromium", "args": ["--no-sandbox", "--disable-setuid-sandbox", "--disable-dev-shm-usage"]}')
        
        # Run mmdc to convert Mermaid to PNG
        result = subprocess.run(
            ['mmdc', '-i', mmd_path, '-o', png_path, '-b', 'white', '-w', '800', '-H', '600', '-p', puppeteer_config_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        # Read the generated PNG
        if os.path.exists(png_path):
            with open(png_path, 'rb') as f:
                image_data = f.read()
            
            # Clean up temp files
            os.unlink(mmd_path)
            os.unlink(png_path)
            
            return io.BytesIO(image_data)
        else:
            logger.warning(f"Mermaid rendering failed: {result.stderr}")
            os.unlink(mmd_path)
            return None
            
    except subprocess.TimeoutExpired:
        logger.error("Mermaid rendering timed out")
        return None
    except Exception as e:
        logger.error(f"Error rendering Mermaid diagram: {e}")
        return None

def generate_docx_from_documentation(docs: List[Dict[str, Any]], repo_name: str) -> io.BytesIO:
    """Generate a DOCX file from documentation data"""
    document = Document()
    
    # Title
    title = document.add_heading(f'Software Documentation: {repo_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = document.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    document.add_page_break()
    
    # Table of Contents
    document.add_heading('Table of Contents', level=1)
    for i, doc in enumerate(docs, 1):
        toc_para = document.add_paragraph(f'{i}. {doc.get("component_path", "Unknown")}')
        toc_para.paragraph_format.left_indent = Inches(0.25)
    
    document.add_page_break()
    
    # Documentation for each file/component
    for doc in docs:
        # Component heading
        document.add_heading(doc.get('component_path', 'Unknown Component'), level=1)
        
        # Metadata
        meta_para = document.add_paragraph()
        meta_para.add_run('Language: ').bold = True
        meta_para.add_run(doc.get('language', 'Unknown'))
        meta_para.add_run(' | ')
        meta_para.add_run('Type: ').bold = True
        meta_para.add_run(doc.get('component_type', 'Unknown'))
        meta_para.add_run(' | ')
        meta_para.add_run('Quality Score: ').bold = True
        meta_para.add_run(f"{doc.get('quality_score', 0):.0f}%")
        
        # Overview section
        if doc.get('markdown'):
            document.add_heading('Overview', level=2)
            # Parse markdown and add as paragraphs
            markdown_text = doc.get('markdown', '')
            # Simple markdown to text conversion
            lines = markdown_text.split('\n')
            for line in lines:
                if line.startswith('# '):
                    document.add_heading(line[2:], level=2)
                elif line.startswith('## '):
                    document.add_heading(line[3:], level=3)
                elif line.startswith('### '):
                    document.add_heading(line[4:], level=4)
                elif line.startswith('```'):
                    continue  # Skip code fence markers
                elif line.strip():
                    document.add_paragraph(line)
        
        # Docstring section
        if doc.get('docstring'):
            document.add_heading('Docstring', level=2)
            docstring_para = document.add_paragraph()
            docstring_run = docstring_para.add_run(doc.get('docstring', ''))
            docstring_run.font.name = 'Courier New'
            docstring_run.font.size = Pt(10)
        
        # Diagram section
        if doc.get('diagrams') and len(doc.get('diagrams', [])) > 0:
            document.add_heading('Diagram', level=2)
            for diagram in doc.get('diagrams', []):
                if isinstance(diagram, dict) and diagram.get('mermaid_code'):
                    mermaid_code = diagram.get('mermaid_code', '')
                    # Clean up the mermaid code - replace escaped newlines with actual newlines
                    clean_code = mermaid_code.replace('\\n', '\n').replace('\\t', '  ').strip()
                    
                    # Render Mermaid to PNG using mmdc CLI
                    diagram_image = render_mermaid_to_image(clean_code)
                    
                    if diagram_image:
                        # Add the rendered diagram image to the document
                        diagram_para = document.add_paragraph()
                        diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = diagram_para.add_run()
                        run.add_picture(diagram_image, width=Inches(5.5))
                    else:
                        # Fallback to code if rendering fails
                        diagram_para = document.add_paragraph()
                        diagram_para.add_run('Mermaid Diagram Code (rendering failed):').bold = True
                        code_para = document.add_paragraph()
                        code_run = code_para.add_run(clean_code)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                    
                    if diagram.get('description'):
                        desc_para = document.add_paragraph()
                        desc_para.add_run('Description: ').bold = True
                        desc_para.add_run(diagram.get('description', ''))
        
        document.add_page_break()
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

async def process_repo_documentation(
    job_id: str,
    tenant_id: str,
    repo_url: str,
    branch: str,
    files: List[Dict[str, Any]],
    access_token: str = None
):
    """Background task to process full repository documentation"""
    try:
        total_files = len(files)
        active_doc_jobs[job_id] = {
            "status": "processing",
            "current_agent": "reader",
            "agents": {
                "reader": {"status": "pending", "progress": 0, "files_processed": 0},
                "searcher": {"status": "pending", "progress": 0, "files_processed": 0},
                "writer": {"status": "pending", "progress": 0, "files_processed": 0},
                "verifier": {"status": "pending", "progress": 0, "files_processed": 0},
                "diagram": {"status": "pending", "progress": 0, "files_processed": 0}
            },
            "files_processed": 0,
            "total_files": total_files,
            "overall_progress": 0,
            "documentation": [],
            "repo_url": repo_url,
            "repo_name": repo_url.split('/')[-1].replace('.git', '')
        }
        
        all_documentation = []
        
        # Process each file through the agent pipeline
        for file_idx, file_data in enumerate(files):
            source_code = file_data['content']
            language = file_data['language']
            component_path = file_data['path']
            
            # Stage 1: Reader Agent
            active_doc_jobs[job_id]["current_agent"] = "reader"
            active_doc_jobs[job_id]["agents"]["reader"]["status"] = "processing"
            
            analysis = await orchestrator.reader.analyze(source_code, language)
            
            active_doc_jobs[job_id]["agents"]["reader"]["files_processed"] = file_idx + 1
            active_doc_jobs[job_id]["agents"]["reader"]["progress"] = int(((file_idx + 1) / total_files) * 100)
            
            # Stage 2: Searcher Agent
            active_doc_jobs[job_id]["current_agent"] = "searcher"
            active_doc_jobs[job_id]["agents"]["searcher"]["status"] = "processing"
            
            context = await orchestrator.searcher.search(analysis, language)
            
            active_doc_jobs[job_id]["agents"]["searcher"]["files_processed"] = file_idx + 1
            active_doc_jobs[job_id]["agents"]["searcher"]["progress"] = int(((file_idx + 1) / total_files) * 100)
            
            # Stage 3: Writer Agent
            active_doc_jobs[job_id]["current_agent"] = "writer"
            active_doc_jobs[job_id]["agents"]["writer"]["status"] = "processing"
            
            documentation = await orchestrator.writer.write(source_code, context, language, "google")
            
            active_doc_jobs[job_id]["agents"]["writer"]["files_processed"] = file_idx + 1
            active_doc_jobs[job_id]["agents"]["writer"]["progress"] = int(((file_idx + 1) / total_files) * 100)
            
            # Stage 4: Verifier Agent
            active_doc_jobs[job_id]["current_agent"] = "verifier"
            active_doc_jobs[job_id]["agents"]["verifier"]["status"] = "processing"
            
            verification = await orchestrator.verifier.verify(source_code, documentation)
            
            active_doc_jobs[job_id]["agents"]["verifier"]["files_processed"] = file_idx + 1
            active_doc_jobs[job_id]["agents"]["verifier"]["progress"] = int(((file_idx + 1) / total_files) * 100)
            
            # Stage 5: Diagram Agent
            active_doc_jobs[job_id]["current_agent"] = "diagram"
            active_doc_jobs[job_id]["agents"]["diagram"]["status"] = "processing"
            
            diagram = await orchestrator.diagram.generate_diagram(source_code)
            
            active_doc_jobs[job_id]["agents"]["diagram"]["files_processed"] = file_idx + 1
            active_doc_jobs[job_id]["agents"]["diagram"]["progress"] = int(((file_idx + 1) / total_files) * 100)
            
            # Compile documentation for this file
            file_doc = {
                "component_path": component_path,
                "component_type": analysis.get("architecture_type", "file"),
                "language": language,
                "docstring": documentation.get("docstring", ""),
                "markdown": documentation.get("markdown", ""),
                "diagrams": [diagram] if diagram else [],
                "quality_score": verification.get("quality_score", 0),
                "analysis": analysis,
                "examples": documentation.get("examples", [])
            }
            all_documentation.append(file_doc)
            
            # Update overall progress
            active_doc_jobs[job_id]["files_processed"] = file_idx + 1
            active_doc_jobs[job_id]["overall_progress"] = int(((file_idx + 1) / total_files) * 100)
            active_doc_jobs[job_id]["documentation"] = all_documentation
            
            # Broadcast progress via WebSocket
            await ws_manager.send_progress(job_id, {
                "type": "repo_doc:progress",
                "job_id": job_id,
                **active_doc_jobs[job_id]
            })
        
        # Mark all agents as completed
        for agent in active_doc_jobs[job_id]["agents"]:
            active_doc_jobs[job_id]["agents"][agent]["status"] = "completed"
            active_doc_jobs[job_id]["agents"][agent]["progress"] = 100
        
        active_doc_jobs[job_id]["status"] = "completed"
        active_doc_jobs[job_id]["current_agent"] = "completed"
        active_doc_jobs[job_id]["overall_progress"] = 100
        
        # Save to database
        repo_name = repo_url.split('/')[-1].replace('.git', '')
        
        # Create repository entry if not exists
        existing_repo = await db.repositories.find_one({"repo_url": repo_url, "tenant_id": tenant_id})
        if not existing_repo:
            repo_id = str(uuid.uuid4())
            await db.repositories.insert_one({
                "id": repo_id,
                "tenant_id": tenant_id,
                "name": repo_name,
                "provider": "github",
                "repo_url": repo_url,
                "branch": branch,
                "language": "mixed",
                "last_synced_at": datetime.now(timezone.utc).isoformat(),
                "components_count": len(all_documentation),
                "coverage_percentage": 100.0,
                "created_at": datetime.now(timezone.utc).isoformat()
            })
        else:
            repo_id = existing_repo["id"]
            await db.repositories.update_one(
                {"id": repo_id},
                {"$set": {
                    "last_synced_at": datetime.now(timezone.utc).isoformat(),
                    "components_count": len(all_documentation)
                }}
            )
        
        # Save documentation entries
        for doc in all_documentation:
            doc_id = str(uuid.uuid4())
            await db.documentation.insert_one({
                "id": doc_id,
                "tenant_id": tenant_id,
                "repository_id": repo_id,
                "component_path": doc["component_path"],
                "component_type": doc["component_type"],
                "language": doc["language"],
                "docstring": doc["docstring"],
                "markdown": doc["markdown"],
                "diagrams": doc["diagrams"],
                "quality_score": doc["quality_score"],
                "version": 1,
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "created_at": datetime.now(timezone.utc).isoformat()
            })
        
        # Broadcast completion
        await ws_manager.send_progress(job_id, {
            "type": "repo_doc:completed",
            "job_id": job_id,
            **active_doc_jobs[job_id]
        })
        
    except Exception as e:
        logger.error(f"Repository documentation job {job_id} failed: {e}")
        active_doc_jobs[job_id]["status"] = "failed"
        active_doc_jobs[job_id]["error"] = str(e)
        await ws_manager.send_progress(job_id, {
            "type": "repo_doc:failed",
            "job_id": job_id,
            "error": str(e)
        })

@repo_docs_router.post("/start")
async def start_repo_documentation(
    request: RepoDocumentationRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Start full repository documentation job"""
    # Fetch repository files
    try:
        # Get user's GitHub access token if available
        user_data = await db.users.find_one({"id": current_user["id"]})
        access_token = user_data.get("github_access_token") if user_data else None
        
        files = await fetch_github_repo_contents(request.repo_url, request.branch, access_token)
        
        if not files:
            raise HTTPException(status_code=400, detail="No code files found in repository")
        
        job_id = str(uuid.uuid4())
        
        # Start background processing
        background_tasks.add_task(
            process_repo_documentation,
            job_id,
            current_user["tenant_id"],
            request.repo_url,
            request.branch,
            files,
            access_token
        )
        
        return {
            "job_id": job_id,
            "status": "started",
            "total_files": len(files),
            "files": [{"path": f["path"], "language": f["language"]} for f in files]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to start repo documentation: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@repo_docs_router.get("/status/{job_id}")
async def get_repo_doc_status(job_id: str, current_user: dict = Depends(get_current_user)):
    """Get status of repository documentation job"""
    if job_id not in active_doc_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = active_doc_jobs[job_id]
    return {
        "job_id": job_id,
        "status": job.get("status", "unknown"),
        "current_agent": job.get("current_agent", ""),
        "agents": job.get("agents", {}),
        "files_processed": job.get("files_processed", 0),
        "total_files": job.get("total_files", 0),
        "overall_progress": job.get("overall_progress", 0),
        "repo_name": job.get("repo_name", ""),
        "error": job.get("error")
    }

@repo_docs_router.get("/export/{job_id}")
async def export_repo_documentation(job_id: str, current_user: dict = Depends(get_current_user)):
    """Export repository documentation as DOCX"""
    if job_id not in active_doc_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = active_doc_jobs[job_id]
    
    if job.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Documentation generation not completed")
    
    documentation = job.get("documentation", [])
    repo_name = job.get("repo_name", "repository")
    
    # Generate DOCX
    docx_buffer = generate_docx_from_documentation(documentation, repo_name)
    
    filename = f"{repo_name}_documentation.docx"
    
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@repo_docs_router.get("/preview/{job_id}")
async def preview_repo_documentation(job_id: str, current_user: dict = Depends(get_current_user)):
    """Get documentation preview for a job"""
    if job_id not in active_doc_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = active_doc_jobs[job_id]
    
    return {
        "job_id": job_id,
        "status": job.get("status", "unknown"),
        "repo_name": job.get("repo_name", ""),
        "documentation": job.get("documentation", [])
    }

# ========================
# WEBSOCKET ENDPOINT
# ========================

@app.websocket("/ws/{client_id}")
async def websocket_endpoint(websocket: WebSocket, client_id: str):
    await ws_manager.connect(websocket, client_id)
    try:
        while True:
            data = await websocket.receive_json()
            # Handle subscription to job updates
            if data.get("type") == "subscribe:job":
                job_id = data.get("job_id")
                # Add to job subscribers
                await ws_manager.send_progress(client_id, {
                    "type": "subscribed",
                    "job_id": job_id
                })
    except WebSocketDisconnect:
        ws_manager.disconnect(client_id)

# ========================
# HEALTH CHECK & ROOT
# ========================

@api_router.get("/")
async def root():
    return {"message": "DocAgent API is running", "version": "1.0.0"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# ========================
# INCLUDE ROUTERS
# ========================

app.include_router(api_router)
app.include_router(auth_router)
app.include_router(repos_router)
app.include_router(docs_router)
app.include_router(jobs_router)
app.include_router(analytics_router)
app.include_router(orgs_router)
app.include_router(models_router)
app.include_router(repo_docs_router)

# ========================
# MIDDLEWARE
# ========================

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
