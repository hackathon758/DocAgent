import os
import re
import io
import logging
import shutil
import subprocess
import sys
import tempfile
import markdown as md_lib
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.mermaid_utils import clean_mermaid_code, validate_mermaid_syntax, attempt_mermaid_repair

logger = logging.getLogger(__name__)


def _find_chrome_executable() -> Optional[str]:
    """Find Chrome/Chromium executable in a platform-aware manner."""
    if sys.platform == "linux":
        candidates = ["/usr/bin/chromium", "/usr/bin/chromium-browser", "/usr/bin/google-chrome"]
    elif sys.platform == "darwin":
        candidates = [
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
            "/Applications/Chromium.app/Contents/MacOS/Chromium",
        ]
    elif sys.platform == "win32":
        candidates = [
            os.path.expandvars(r"%ProgramFiles%\Google\Chrome\Application\chrome.exe"),
            os.path.expandvars(r"%ProgramFiles(x86)%\Google\Chrome\Application\chrome.exe"),
            os.path.expandvars(r"%LocalAppData%\Google\Chrome\Application\chrome.exe"),
        ]
    else:
        candidates = []

    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def render_mermaid_to_image(mermaid_code: str) -> Optional[io.BytesIO]:
    """Render Mermaid code to PNG image using mmdc CLI (cross-platform)."""
    mmdc_path = shutil.which("mmdc")
    if not mmdc_path:
        logger.warning("mmdc (Mermaid CLI) not found on PATH — cannot render diagram to image")
        return None

    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as mmd_file:
            mmd_file.write(mermaid_code)
            mmd_path = mmd_file.name

        png_path = mmd_path.replace('.mmd', '.png')

        # Build puppeteer config with platform-aware Chrome path
        puppeteer_config_path = os.path.join(tempfile.gettempdir(), 'puppeteer-config.json')
        if not os.path.exists(puppeteer_config_path):
            chrome_path = _find_chrome_executable()
            config = {"args": ["--no-sandbox", "--disable-setuid-sandbox", "--disable-dev-shm-usage"]}
            if chrome_path:
                config["executablePath"] = chrome_path
            import json
            with open(puppeteer_config_path, 'w') as f:
                json.dump(config, f)

        cmd = [mmdc_path, '-i', mmd_path, '-o', png_path, '-b', 'white', '-w', '800', '-H', '600', '-p', puppeteer_config_path]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        if os.path.exists(png_path) and os.path.getsize(png_path) > 0:
            with open(png_path, 'rb') as f:
                image_data = f.read()
            os.unlink(mmd_path)
            os.unlink(png_path)
            return io.BytesIO(image_data)
        else:
            logger.warning(f"Mermaid rendering failed (mmdc exit {result.returncode}): {result.stderr[:300]}")
            if os.path.exists(mmd_path):
                os.unlink(mmd_path)
            if os.path.exists(png_path):
                os.unlink(png_path)
            return None

    except subprocess.TimeoutExpired:
        logger.error("Mermaid rendering timed out")
        return None
    except Exception as e:
        logger.error(f"Error rendering Mermaid diagram: {e}")
        return None


def generate_docx_from_documentation(docs: List[Dict[str, Any]], repo_name: str) -> io.BytesIO:
    """Generate a DOCX file from documentation data"""
    document = Document()

    title = document.add_heading(f'Software Documentation: {repo_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = document.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)

    document.add_page_break()

    document.add_heading('Table of Contents', level=1)
    for i, doc in enumerate(docs, 1):
        toc_para = document.add_paragraph(f'{i}. {doc.get("component_path", "Unknown")}')
        toc_para.paragraph_format.left_indent = Inches(0.25)

    document.add_page_break()

    for doc in docs:
        document.add_heading(doc.get('component_path', 'Unknown Component'), level=1)

        meta_para = document.add_paragraph()
        meta_para.add_run('Language: ').bold = True
        meta_para.add_run(doc.get('language', 'Unknown'))
        meta_para.add_run(' | ')
        meta_para.add_run('Type: ').bold = True
        meta_para.add_run(doc.get('component_type', 'Unknown'))
        meta_para.add_run(' | ')
        meta_para.add_run('Quality Score: ').bold = True
        meta_para.add_run(f"{doc.get('quality_score', 0):.0f}%")

        if doc.get('markdown'):
            document.add_heading('Overview', level=2)
            markdown_text = doc.get('markdown', '')
            lines = markdown_text.split('\n')
            for line in lines:
                if line.startswith('# '):
                    document.add_heading(line[2:], level=2)
                elif line.startswith('## '):
                    document.add_heading(line[3:], level=3)
                elif line.startswith('### '):
                    document.add_heading(line[4:], level=4)
                elif line.startswith('```'):
                    continue
                elif line.strip():
                    document.add_paragraph(line)

        if doc.get('docstring'):
            document.add_heading('Docstring', level=2)
            docstring_para = document.add_paragraph()
            docstring_run = docstring_para.add_run(doc.get('docstring', ''))
            docstring_run.font.name = 'Courier New'
            docstring_run.font.size = Pt(10)

        if doc.get('diagrams') and len(doc.get('diagrams', [])) > 0:
            document.add_heading('Diagram', level=2)
            for diagram in doc.get('diagrams', []):
                if isinstance(diagram, dict) and diagram.get('mermaid_code'):
                    mermaid_code = diagram.get('mermaid_code', '')
                    clean_code = clean_mermaid_code(mermaid_code)

                    if not clean_code:
                        logger.warning(f"Empty diagram code after cleaning for {doc.get('component_path', '?')}")
                        continue

                    # Validate and attempt repair before rendering
                    is_valid, err = validate_mermaid_syntax(clean_code)
                    if not is_valid:
                        logger.warning(f"Invalid Mermaid syntax for {doc.get('component_path', '?')}: {err} — attempting repair")
                        clean_code = attempt_mermaid_repair(clean_code, err)
                        is_valid2, err2 = validate_mermaid_syntax(clean_code)
                        if not is_valid2:
                            logger.warning(f"Repair failed for {doc.get('component_path', '?')}: {err2} — embedding as raw code")

                    # Try rendering to image
                    diagram_image = render_mermaid_to_image(clean_code)

                    # Retry with repaired code if first attempt failed
                    if not diagram_image and is_valid:
                        repaired = attempt_mermaid_repair(clean_code, "render failure")
                        if repaired != clean_code:
                            logger.info("Retrying render with repaired code")
                            diagram_image = render_mermaid_to_image(repaired)

                    if diagram_image:
                        diagram_para = document.add_paragraph()
                        diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = diagram_para.add_run()
                        run.add_picture(diagram_image, width=Inches(5.5))
                    else:
                        diagram_para = document.add_paragraph()
                        diagram_para.add_run('Mermaid Diagram Code (rendering failed):').bold = True
                        code_para = document.add_paragraph()
                        code_run = code_para.add_run(clean_code)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)

                    if diagram.get('description'):
                        desc_para = document.add_paragraph()
                        desc_para.add_run('Description: ').bold = True
                        desc_para.add_run(diagram.get('description', ''))

        document.add_page_break()

    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


# ---------------------------------------------------------------------------
# Helpers for comprehensive doc
# ---------------------------------------------------------------------------

def _add_info_row(document, label: str, value: str):
    p = document.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(str(value))


def _add_placeholder(document, text: str):
    p = document.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.font.color.rgb = RGBColor(160, 160, 160)
    run.font.italic = True
    run.font.size = Pt(10)


def _add_code_block(document, code: str, font_size: int = 9):
    p = document.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(font_size)


def _detect_project_metadata(docs, file_contents, repo_name, repo_url, branch):
    """Extract project-level metadata from per-file results and raw content."""
    import re as _re

    languages = {}
    total_quality = 0
    quality_count = 0
    all_diagrams = []
    api_endpoints = []
    db_models = []
    config_files = []
    test_files_detected = []
    dependency_files = {}
    security_patterns = []
    all_files = []

    route_pats = [
        _re.compile(r'@(?:app|router|blueprint)\.(get|post|put|delete|patch)\s*\(\s*["\']([^"\']+)', _re.I),
        _re.compile(r'router\.(get|post|put|delete|patch)\s*\(\s*["\']([^"\']+)', _re.I),
    ]
    model_pats = [
        _re.compile(r'class\s+(\w+)\s*\(.*(?:Model|Base|Document|Schema|BaseModel)', _re.I),
        _re.compile(r'CREATE\s+TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?[`\'"]?(\w+)', _re.I),
    ]
    sec_kw = _re.compile(r'(auth|token|jwt|oauth|password|encrypt|secret|api_key|csrf|cors|ssl|permission|role)', _re.I)

    dep_names = {"package.json", "requirements.txt", "Pipfile", "pyproject.toml", "Gemfile", "go.mod", "Cargo.toml", "pom.xml", "composer.json"}
    cfg_names = {".env", ".env.example", "config.py", "config.js", "settings.py", "docker-compose.yml", "Dockerfile", "tsconfig.json"}

    for doc in docs:
        path = doc.get("component_path", "")
        lang = doc.get("language", "text")
        languages[lang] = languages.get(lang, 0) + 1
        all_files.append(path)
        if doc.get("quality_score"):
            total_quality += doc["quality_score"]
            quality_count += 1
        if doc.get("diagram"):
            all_diagrams.append({"code": doc["diagram"], "source": path})

    for path, content in file_contents.items():
        basename = path.replace("\\", "/").split("/")[-1]
        if basename.lower() in dep_names:
            dependency_files[basename] = content[:3000]
        if basename.lower() in cfg_names or basename.startswith(".env"):
            config_files.append(path)
        if any(t in path.lower() for t in ("test_", "_test.", ".test.", "spec.", "tests/")):
            test_files_detected.append(path)
        for pat in route_pats:
            for m in pat.finditer(content):
                g = m.groups()
                if len(g) == 2:
                    api_endpoints.append({"method": g[0].upper(), "path": g[1], "file": path})
        for pat in model_pats:
            for m in pat.finditer(content):
                db_models.append({"name": m.group(1), "file": path})
        for m in sec_kw.finditer(content):
            kw = m.group(1).lower()
            if kw not in [s["keyword"] for s in security_patterns]:
                security_patterns.append({"keyword": kw, "file": path})

    avg_quality = round(total_quality / quality_count) if quality_count else 0
    return {
        "repo_name": repo_name, "repo_url": repo_url, "branch": branch,
        "languages": languages, "avg_quality": avg_quality, "total_files": len(all_files),
        "all_files": all_files, "all_diagrams": all_diagrams, "api_endpoints": api_endpoints,
        "db_models": db_models, "config_files": config_files, "test_files": test_files_detected,
        "dependency_files": dependency_files, "security_patterns": security_patterns,
    }


# ---------------------------------------------------------------------------
# Comprehensive Software Delivery Document Generator
# ---------------------------------------------------------------------------

def generate_comprehensive_docx(
    docs: List[Dict[str, Any]],
    repo_name: str,
    repo_url: str = "",
    branch: str = "main",
    file_contents: Optional[Dict[str, str]] = None,
) -> io.BytesIO:
    """Generate a comprehensive 20-section software delivery DOCX document."""
    if file_contents is None:
        file_contents = {}

    meta = _detect_project_metadata(docs, file_contents, repo_name, repo_url, branch)
    document = Document()
    now_str = datetime.now().strftime("%B %d, %Y")
    lang_summary = ", ".join(f"{l} ({c})" for l, c in sorted(meta["languages"].items(), key=lambda x: -x[1]))

    # ── COVER PAGE ──
    for _ in range(6):
        document.add_paragraph("")
    t = document.add_heading("Software Documentation", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = document.add_heading(repo_name, level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp = document.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"Generated on {now_str}")
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(128, 128, 128)
    vp = document.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("Version 1.0 | Initial Release")
    vr.font.size = Pt(11)
    vr.font.color.rgb = RGBColor(128, 128, 128)
    document.add_page_break()

    # ── TABLE OF CONTENTS ──
    document.add_heading("Table of Contents", level=1)
    sections = [
        "1. Project Information", "2. Executive Summary", "3. Scope of Delivery",
        "4. System Requirements", "5. Installation Guide", "6. System Architecture",
        "7. Database Schema", "8. API Documentation", "9. Code Documentation",
        "10. User Manual", "11. Admin Guide", "12. Source Code Delivery",
        "13. Test Documentation", "14. Training Materials", "15. Release Notes",
        "16. Support & Maintenance", "17. Security Documentation",
        "18. Post-Deployment Checklist", "19. Sign-Off & Acceptance", "20. Appendices",
    ]
    for item in sections:
        p = document.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)
    document.add_page_break()

    # ── 1. PROJECT INFORMATION ──
    document.add_heading("1. Project Information", level=1)
    _add_info_row(document, "Project Name", repo_name)
    _add_info_row(document, "Repository URL", repo_url or "N/A")
    _add_info_row(document, "Branch", branch)
    _add_info_row(document, "Delivery Date", now_str)
    _add_info_row(document, "Version", "1.0")
    _add_info_row(document, "Release Type", "Initial Release")
    _add_placeholder(document, "Client Name — enter the client or stakeholder name")
    _add_placeholder(document, "Project ID — enter internal project identifier")
    document.add_page_break()

    # ── 2. EXECUTIVE SUMMARY ──
    document.add_heading("2. Executive Summary", level=1)
    document.add_heading("Project Overview", level=2)
    document.add_paragraph(
        f"This document provides comprehensive software documentation for the "
        f"{repo_name} project. The codebase consists of {meta['total_files']} source files "
        f"spanning: {lang_summary}."
    )
    document.add_heading("Objectives Achieved", level=2)
    document.add_paragraph(
        f"Automated documentation completed with average quality score of {meta['avg_quality']}%. "
        f"Total of {len(meta['all_diagrams'])} architectural diagrams generated."
    )
    document.add_heading("Key Deliverables", level=2)
    for d in [
        f"Complete code documentation for {meta['total_files']} files",
        f"{len(meta['api_endpoints'])} API endpoints documented" if meta["api_endpoints"] else "API endpoint inventory",
        f"{len(meta['db_models'])} database models/schemas identified" if meta["db_models"] else "Database schema analysis",
        f"{len(meta['all_diagrams'])} architectural diagrams",
        "Quality verification reports",
    ]:
        document.add_paragraph(d, style="List Bullet")
    document.add_page_break()

    # ── 3. SCOPE OF DELIVERY ──
    document.add_heading("3. Scope of Delivery", level=1)
    document.add_heading("Modules Included", level=2)
    modules = {}
    for f in meta["all_files"]:
        parts = f.replace("\\", "/").split("/")
        mod = parts[0] if len(parts) > 1 else "(root)"
        modules.setdefault(mod, []).append(f)
    for mod_name, mod_files in sorted(modules.items()):
        document.add_paragraph(f"{mod_name} ({len(mod_files)} files)", style="List Bullet")
    document.add_heading("Excluded Items", level=2)
    for d in ["node_modules", "venv", ".git", "__pycache__", "dist", "build", "vendor", "coverage"]:
        document.add_paragraph(d, style="List Bullet")
    document.add_heading("Acceptance Criteria", level=2)
    _add_placeholder(document, "Define acceptance criteria and sign-off requirements here")
    document.add_page_break()

    # ── 4. SYSTEM REQUIREMENTS ──
    document.add_heading("4. System Requirements", level=1)
    document.add_heading("Technology Stack", level=2)
    _add_info_row(document, "Languages", lang_summary)
    document.add_heading("Third-Party Dependencies", level=2)
    if meta["dependency_files"]:
        for dep_file, content in meta["dependency_files"].items():
            document.add_heading(dep_file, level=3)
            lines = content.split("\n")[:60]
            _add_code_block(document, "\n".join(lines))
    else:
        _add_placeholder(document, "No dependency files detected. List prerequisites here.")
    document.add_heading("Hardware & OS Compatibility", level=2)
    _add_placeholder(document, "Specify hardware specs, OS compatibility, and browser support")
    document.add_page_break()

    # ── 5. INSTALLATION GUIDE ──
    document.add_heading("5. Installation Guide", level=1)
    if meta["config_files"]:
        document.add_heading("Configuration Files Detected", level=2)
        for cf in meta["config_files"]:
            document.add_paragraph(cf, style="List Bullet")
    document.add_heading("Deployment Steps", level=2)
    _add_placeholder(document, "Provide step-by-step deployment instructions")
    document.add_heading("Database Setup", level=2)
    _add_placeholder(document, "Include database setup scripts and migration procedures")
    document.add_page_break()

    # ── 6. SYSTEM ARCHITECTURE ──
    document.add_heading("6. System Architecture", level=1)
    if meta["all_diagrams"]:
        document.add_paragraph(f"{len(meta['all_diagrams'])} diagrams generated from the codebase:")
        for i, diag in enumerate(meta["all_diagrams"][:10]):
            document.add_heading(f"Diagram: {diag.get('source', f'Component {i+1}')}", level=3)
            code = diag.get("code", "")
            if code:
                cc = clean_mermaid_code(code)
                if cc:
                    img = render_mermaid_to_image(cc)
                    if img:
                        ip = document.add_paragraph()
                        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        ip.add_run().add_picture(img, width=Inches(5.5))
                    else:
                        document.add_paragraph("Mermaid Diagram Code:").runs[0].bold = True
                        _add_code_block(document, cc)
    else:
        _add_placeholder(document, "Insert high-level architecture diagram here")
    document.add_heading("Design Decisions", level=2)
    _add_placeholder(document, "Document key design decisions and rationale")
    document.add_heading("Integration Points", level=2)
    _add_placeholder(document, "List external integrations and dependencies")
    document.add_page_break()

    # ── 7. DATABASE SCHEMA ──
    document.add_heading("7. Database Schema", level=1)
    if meta["db_models"]:
        document.add_heading("Detected Models / Entities", level=2)
        tbl = document.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Model / Entity"
        hdr[1].text = "Source File"
        for hc in hdr:
            for p in hc.paragraphs:
                for r in p.runs:
                    r.bold = True
        seen = set()
        for model in meta["db_models"]:
            if model["name"] not in seen:
                seen.add(model["name"])
                row = tbl.add_row().cells
                row[0].text = model["name"]
                row[1].text = model["file"]
    else:
        _add_placeholder(document, "No database models/schemas were automatically detected")
    document.add_heading("Entity Relationships", level=2)
    _add_placeholder(document, "Insert entity-relationship diagrams and table definitions")
    document.add_page_break()

    # ── 8. API DOCUMENTATION ──
    document.add_heading("8. API Documentation", level=1)
    if meta["api_endpoints"]:
        document.add_paragraph(f"{len(meta['api_endpoints'])} API endpoints detected:")
        tbl = document.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Method"
        hdr[1].text = "Endpoint"
        hdr[2].text = "Source File"
        for hc in hdr:
            for p in hc.paragraphs:
                for r in p.runs:
                    r.bold = True
        seen_ep = set()
        for ep in meta["api_endpoints"]:
            key = f"{ep['method']} {ep['path']}"
            if key not in seen_ep:
                seen_ep.add(key)
                row = tbl.add_row().cells
                row[0].text = ep["method"]
                row[1].text = ep["path"]
                row[2].text = ep["file"]
        document.add_heading("Authentication", level=2)
        auth_eps = [ep for ep in meta["api_endpoints"] if "auth" in ep["path"].lower()]
        if auth_eps:
            for ep in auth_eps:
                document.add_paragraph(f"{ep['method']} {ep['path']}", style="List Bullet")
        else:
            _add_placeholder(document, "Document authentication methods here")
    else:
        _add_placeholder(document, "No API endpoints detected. Document endpoints manually.")
    document.add_heading("Error Codes", level=2)
    _add_placeholder(document, "List error codes, formats, and response examples")
    document.add_page_break()

    # ── 9. CODE DOCUMENTATION ──
    document.add_heading("9. Code Documentation", level=1)
    document.add_paragraph(
        f"Detailed documentation for {len(docs)} source files, generated through the "
        f"5-agent AI pipeline (Reader → Searcher → Writer → Verifier → Diagram)."
    )
    code_mods = {}
    for doc in docs:
        path = doc.get("component_path", "unknown")
        parts = path.replace("\\", "/").split("/")
        mod = parts[0] if len(parts) > 1 else "(root)"
        code_mods.setdefault(mod, []).append(doc)
    for mod_name, mod_docs in sorted(code_mods.items()):
        document.add_heading(f"Module: {mod_name}", level=2)
        for doc in mod_docs:
            document.add_heading(doc.get("component_path", "Unknown"), level=3)
            mp = document.add_paragraph()
            mp.add_run("Language: ").bold = True
            mp.add_run(doc.get("language", "Unknown"))
            mp.add_run(" | ")
            mp.add_run("Quality: ").bold = True
            mp.add_run(f"{doc.get('quality_score', 0):.0f}%")
            if doc.get("markdown"):
                for line in doc["markdown"].split("\n"):
                    if line.startswith("# "):
                        document.add_heading(line[2:], level=4)
                    elif line.startswith("## "):
                        document.add_heading(line[3:], level=4)
                    elif line.startswith("```"):
                        continue
                    elif line.strip():
                        document.add_paragraph(line)
            if doc.get("docstring"):
                document.add_heading("Docstring", level=4)
                dp2 = document.add_paragraph()
                rn = dp2.add_run(doc["docstring"])
                rn.font.name = "Courier New"
                rn.font.size = Pt(9)
            if doc.get("usage_example"):
                document.add_heading("Usage Example", level=4)
                _add_code_block(document, doc["usage_example"])
            if doc.get("complexity"):
                _add_info_row(document, "Complexity", doc["complexity"])
    document.add_page_break()

    # ── 10. USER MANUAL ──
    document.add_heading("10. User Manual", level=1)
    document.add_heading("Feature Walkthroughs", level=2)
    _add_placeholder(document, "Provide step-by-step walkthroughs for each major feature")
    document.add_heading("Use Cases", level=2)
    _add_placeholder(document, "List primary use cases with screen-by-screen instructions")
    document.add_heading("User Roles & Permissions", level=2)
    _add_placeholder(document, "Define user roles and their permissions matrix")
    document.add_page_break()

    # ── 11. ADMIN GUIDE ──
    document.add_heading("11. Admin Guide", level=1)
    document.add_heading("Administrative Functions", level=2)
    _add_placeholder(document, "Document system administration functions")
    document.add_heading("User Management", level=2)
    _add_placeholder(document, "Explain user management procedures")
    document.add_heading("Backup & Restore", level=2)
    _add_placeholder(document, "Document backup and restore procedures")
    document.add_page_break()

    # ── 12. SOURCE CODE DELIVERY ──
    document.add_heading("12. Source Code Delivery", level=1)
    _add_info_row(document, "Repository URL", repo_url or "N/A")
    _add_info_row(document, "Branch", branch)
    document.add_heading("Folder Structure", level=2)
    tree_lines = []
    sorted_files = sorted(meta["all_files"])
    for fp in sorted_files[:80]:
        parts = fp.replace("\\", "/").split("/")
        for i, part in enumerate(parts):
            prefix = "  " * i + ("└── " if i == len(parts) - 1 else "├── ")
            tree_lines.append(f"{prefix}{part}")
    if tree_lines:
        _add_code_block(document, "\n".join(tree_lines))
    if len(sorted_files) > 80:
        document.add_paragraph(f"... and {len(sorted_files) - 80} more files")
    document.add_heading("Build Instructions", level=2)
    _add_placeholder(document, "Provide build instructions and deployment scripts")
    document.add_page_break()

    # ── 13. TEST DOCUMENTATION ──
    document.add_heading("13. Test Documentation", level=1)
    if meta["test_files"]:
        document.add_heading("Test Files Detected", level=2)
        for tf in meta["test_files"][:30]:
            document.add_paragraph(tf, style="List Bullet")
    else:
        document.add_paragraph("No test files were automatically detected.")
    document.add_heading("Quality Assurance Summary", level=2)
    _add_info_row(document, "Files Analyzed", str(meta["total_files"]))
    _add_info_row(document, "Average Quality Score", f"{meta['avg_quality']}%")
    _add_info_row(document, "Diagrams Generated", str(len(meta["all_diagrams"])))
    document.add_heading("Test Plan", level=2)
    _add_placeholder(document, "Include test plan, test cases, and results summary")
    document.add_page_break()

    # ── 14. TRAINING MATERIALS ──
    document.add_heading("14. Training Materials", level=1)
    _add_placeholder(document, "Provide user training guides and onboarding materials")
    _add_placeholder(document, "List video tutorials with links")
    _add_placeholder(document, "Include quick reference cards and cheat sheets")
    document.add_page_break()

    # ── 15. RELEASE NOTES ──
    document.add_heading("15. Release Notes", level=1)
    document.add_heading("Version 1.0 — Initial Release", level=2)
    _add_info_row(document, "Release Date", now_str)
    _add_info_row(document, "Total Files", str(meta["total_files"]))
    _add_info_row(document, "Languages", lang_summary)
    document.add_heading("New Features", level=3)
    _add_placeholder(document, "List new features introduced in this release")
    document.add_heading("Known Issues", level=3)
    _add_placeholder(document, "List known issues and workarounds")
    document.add_page_break()

    # ── 16. SUPPORT & MAINTENANCE ──
    document.add_heading("16. Support & Maintenance", level=1)
    document.add_heading("Troubleshooting Guide", level=2)
    _add_placeholder(document, "Common issues and resolution steps")
    document.add_heading("FAQ", level=2)
    _add_placeholder(document, "Frequently asked questions")
    document.add_heading("SLA Terms", level=2)
    _add_placeholder(document, "Service Level Agreement terms and escalation process")
    document.add_page_break()

    # ── 17. SECURITY DOCUMENTATION ──
    document.add_heading("17. Security Documentation", level=1)
    if meta["security_patterns"]:
        document.add_heading("Security Patterns Detected", level=2)
        kw_files = {}
        for sp in meta["security_patterns"]:
            kw_files.setdefault(sp["keyword"], []).append(sp["file"])
        tbl = document.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Security Feature"
        hdr[1].text = "Source File(s)"
        for hc in hdr:
            for p in hc.paragraphs:
                for r in p.runs:
                    r.bold = True
        for kw, fls in sorted(kw_files.items()):
            row = tbl.add_row().cells
            row[0].text = kw.upper()
            row[1].text = ", ".join(set(fls))[:200]
    else:
        document.add_paragraph("No specific security patterns were detected.")
    document.add_heading("Compliance & Certifications", level=2)
    _add_placeholder(document, "List compliance certifications and standards met")
    document.add_heading("Data Privacy", level=2)
    _add_placeholder(document, "Document data privacy policies and access control guidelines")
    document.add_page_break()

    # ── 18. POST-DEPLOYMENT CHECKLIST ──
    document.add_heading("18. Post-Deployment Checklist", level=1)
    checklist = [
        "Environment variables configured", "Database migrations executed",
        "API endpoints responding", "Authentication tested", "SSL/TLS verified",
        "Monitoring configured", "Backup procedures tested", "Load testing completed",
        "Security scan passed", "DNS configured", "Logging enabled", "UAT completed",
    ]
    tbl = document.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Item"
    hdr[2].text = "Status"
    for hc in hdr:
        for p in hc.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, item in enumerate(checklist, 1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = item
        row[2].text = "☐ Pending"
    document.add_heading("Rollback Procedures", level=2)
    _add_placeholder(document, "Document rollback procedures and contingency plans")
    document.add_page_break()

    # ── 19. SIGN-OFF & ACCEPTANCE ──
    document.add_heading("19. Sign-Off & Acceptance", level=1)
    for item in [
        "Source code delivered", "Documentation approved", "Test results accepted",
        "Security assessment completed", "Training delivered", "Knowledge transfer done",
    ]:
        document.add_paragraph(f"☐  {item}")
    document.add_heading("Client Acceptance", level=2)
    for label in ["Accepted By", "Title", "Signature", "Date"]:
        _add_info_row(document, label, "________________________________")
    document.add_paragraph("")
    for label in ["Delivered By", "Title", "Signature", "Date"]:
        _add_info_row(document, label, "________________________________")
    document.add_page_break()

    # ── 20. APPENDICES ──
    document.add_heading("20. Appendices", level=1)
    document.add_heading("A. Glossary of Terms", level=2)
    glossary = {
        "API": "Application Programming Interface", "CI/CD": "Continuous Integration / Continuous Deployment",
        "CORS": "Cross-Origin Resource Sharing", "CRUD": "Create, Read, Update, Delete",
        "JWT": "JSON Web Token", "ORM": "Object-Relational Mapping",
        "REST": "Representational State Transfer", "SLA": "Service Level Agreement",
        "SSL/TLS": "Secure Sockets Layer / Transport Layer Security",
    }
    tbl = document.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Term"
    hdr[1].text = "Definition"
    for hc in hdr:
        for p in hc.paragraphs:
            for r in p.runs:
                r.bold = True
    for term, defn in glossary.items():
        row = tbl.add_row().cells
        row[0].text = term
        row[1].text = defn

    document.add_heading("B. Per-File Details", level=2)
    for i, doc in enumerate(docs, 1):
        p = document.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(f"{doc.get('component_path', 'Unknown')}  [{doc.get('language', '?')}]  Quality: {doc.get('quality_score', 0):.0f}%")

    document.add_heading("C. Version History", level=2)
    tbl = document.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Version"
    hdr[1].text = "Date"
    hdr[2].text = "Author"
    hdr[3].text = "Changes"
    for hc in hdr:
        for p in hc.paragraphs:
            for r in p.runs:
                r.bold = True
    row = tbl.add_row().cells
    row[0].text = "1.0"
    row[1].text = now_str
    row[2].text = "DocAgent AI"
    row[3].text = "Initial documentation generation"

    # ── SAVE ──
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


async def fetch_github_repo_contents(repo_url: str, branch: str = "main", access_token: str = None) -> Dict[str, Any]:
    """Fetch code files, metadata files, and test file info from a GitHub repository.

    Uses `git clone --depth 1` which has no API rate limit for public repos,
    then reads the files from the cloned directory.

    Returns a dict with:
      - files: list of code file dicts (path, name, content, language, size)
      - metadata_files: list of metadata file dicts (path, name, content)
      - test_files: list of test file info dicts (path, name, size)
    """
    import asyncio
    from fastapi import HTTPException

    match = re.match(r'(?:https?://)?github\.com/([^/]+)/([^/]+?)(?:\.git)?/?$', repo_url)
    if not match:
        raise HTTPException(status_code=400, detail="Invalid GitHub repository URL")

    owner, repo_name = match.groups()

    code_extensions = {'.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h', '.go', '.rs', '.cs', '.rb', '.php'}
    test_dirs = {'test', 'tests', '__tests__', 'spec', 'specs'}
    skip_dirs = {'node_modules', 'venv', '.git', '__pycache__', 'dist', 'build', '.idea', '.vscode',
                 'vendor', '.github', 'coverage', '.nyc_output'}
    metadata_filenames = {
        'readme.md', 'readme.rst', 'readme.txt', 'readme',
        'license', 'license.md', 'license.txt',
        'package.json', 'package-lock.json',
        'requirements.txt', 'setup.py', 'setup.cfg', 'pyproject.toml',
        'dockerfile', 'docker-compose.yml', 'docker-compose.yaml',
        '.env.example', '.env.sample',
        'makefile', 'cargo.toml', 'go.mod', 'go.sum',
        'tsconfig.json', 'webpack.config.js', 'vite.config.js', 'vite.config.ts',
        'changelog.md', 'changelog', 'changes.md',
    }
    lang_map = {
        '.py': 'python', '.js': 'javascript', '.jsx': 'javascript',
        '.ts': 'typescript', '.tsx': 'typescript', '.java': 'java',
        '.cpp': 'cpp', '.c': 'c', '.h': 'c', '.go': 'go',
        '.rs': 'rust', '.cs': 'csharp', '.rb': 'ruby', '.php': 'php'
    }

    clone_url = f"https://github.com/{owner}/{repo_name}.git"
    if access_token:
        clone_url = f"https://x-access-token:{access_token}@github.com/{owner}/{repo_name}.git"

    clone_dir = tempfile.mkdtemp(prefix="docagent_")
    # git clone needs a non-existent target path
    import shutil
    shutil.rmtree(clone_dir, ignore_errors=True)
    try:
        # Shallow clone (only latest commit, single branch)
        env = os.environ.copy()
        env["GIT_TERMINAL_PROMPT"] = "0"  # Prevent git from prompting for credentials

        # Try cloning with specified branch first
        proc = await asyncio.create_subprocess_exec(
            "git", "clone", "--depth", "1", "--branch", branch, "--single-branch", clone_url, clone_dir,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            env=env,
        )
        _, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)

        if proc.returncode != 0:
            err_msg = stderr.decode().strip() if stderr else "Unknown error"
            logger.warning(f"git clone with branch '{branch}' failed: {err_msg}")

            # Retry without --branch flag to clone the repo's default branch
            shutil.rmtree(clone_dir, ignore_errors=True)
            logger.info(f"Retrying clone without --branch (using repo default branch)")
            proc2 = await asyncio.create_subprocess_exec(
                "git", "clone", "--depth", "1", clone_url, clone_dir,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
            _, stderr2 = await asyncio.wait_for(proc2.communicate(), timeout=120)

            if proc2.returncode != 0:
                err_msg2 = stderr2.decode().strip() if stderr2 else "Unknown error"
                logger.error(f"git clone (default branch) also failed: {err_msg2}")
                if "not found" in err_msg2.lower() or "404" in err_msg2 or "authentication" in err_msg2.lower():
                    raise HTTPException(status_code=404, detail="Repository not found. Check the URL and ensure the repo is accessible.")
                raise HTTPException(status_code=502, detail=f"Failed to clone repository: {err_msg2}")

        files = []
        metadata_files = []
        test_files = []
        MAX_FILES = 50

        for root, dirs, filenames in os.walk(clone_dir):
            # Remove dirs we never want to recurse into
            dirs[:] = [d for d in dirs if d not in skip_dirs]

            rel_root = os.path.relpath(root, clone_dir).replace("\\", "/")
            # Detect if we're inside a test directory
            root_parts = set(rel_root.split("/"))
            is_test_dir = bool(root_parts & test_dirs)

            for fname in filenames:
                full_path = os.path.join(root, fname)
                rel_path = os.path.relpath(full_path, clone_dir).replace("\\", "/")
                ext = os.path.splitext(fname)[1].lower()

                # Metadata files (root-level or any depth for Dockerfiles etc.)
                if fname.lower() in metadata_filenames:
                    try:
                        with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                            content = f.read()
                        metadata_files.append({
                            'path': rel_path,
                            'name': fname,
                            'content': content,
                        })
                    except Exception as e:
                        logger.warning(f"Could not read metadata file {rel_path}: {e}")
                    continue

                # Test files — collect path/name/size only
                if is_test_dir and ext in code_extensions:
                    try:
                        size = os.path.getsize(full_path)
                        test_files.append({
                            'path': rel_path,
                            'name': fname,
                            'size': size,
                            'language': lang_map.get(ext, 'text'),
                        })
                    except Exception:
                        pass
                    continue

                # Code files
                if ext not in code_extensions:
                    continue
                if len(files) >= MAX_FILES:
                    continue
                try:
                    with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                    files.append({
                        'path': rel_path,
                        'name': fname,
                        'content': content,
                        'language': lang_map.get(ext, 'text'),
                        'size': len(content),
                    })
                except Exception as e:
                    logger.warning(f"Could not read {rel_path}: {e}")

        return {
            "files": files,
            "metadata_files": metadata_files,
            "test_files": test_files,
        }

    except HTTPException:
        raise
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="Repository clone timed out")
    except Exception as e:
        logger.error(f"Error fetching repo contents: {e}")
        raise HTTPException(status_code=502, detail=f"Could not fetch repository contents: {e}")
    finally:
        # Clean up cloned directory
        try:
            shutil.rmtree(clone_dir, ignore_errors=True)
        except Exception:
            pass
